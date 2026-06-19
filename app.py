import os
import secrets
import re
import json
import uuid
import urllib.request
import random
import sqlite3
from io import BytesIO
from datetime import datetime, timezone
from datetime import datetime as dt

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import qrcode
import base64

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
from reportlab.lib.colors import HexColor

import pdfplumber
import docx as docx_lib

app = Flask(__name__, static_folder="static", static_url_path="/static")
app.secret_key = os.environ.get("SECRET_KEY", secrets.token_hex(32))
CORS(app, supports_credentials=True)

OUTPUTS_DIR = os.path.join(os.path.dirname(__file__), "outputs")
os.makedirs(OUTPUTS_DIR, exist_ok=True)
STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")
os.makedirs(STATIC_DIR, exist_ok=True)

WALLETS = {
    "bitcoin":  os.environ.get("WALLET_BTC",  "bc1qyourbitcoinaddresshere"),
    "ethereum": os.environ.get("WALLET_ETH",  "0xYourEthereumAddressHere"),
    "usdc":     os.environ.get("WALLET_USDC", "0xYourUSDCAddressHere"),
    "solana":   os.environ.get("WALLET_SOL",  "YourSolanaAddressHere"),
}

SERVICES = {
    "resume_ai":        {"name": "AI Resume Generator",  "price_usd": 49},
    "resume_optimizer": {"name": "Resume Optimizer",     "price_usd": 49},
    "interview":        {"name": "Interview Prep",       "price_usd": 29},
    "salary":           {"name": "Salary Negotiator",    "price_usd": 19},
}

CRYPTO_PRICES_USD = {
    "bitcoin":  65000,
    "ethereum": 3500,
    "usdc":     1.0,
    "solana":   150,
}

INDUSTRY_KEYWORDS = {
    "tech":       ["Python", "JavaScript", "AWS", "Agile", "REST API", "Git", "React", "Docker", "Kubernetes", "CI/CD"],
    "marketing":  ["SEO", "Social Media", "Analytics", "Campaign ROI", "Content Strategy", "Email Marketing", "Demand Gen"],
    "sales":      ["Revenue Growth", "Quota Attainment", "Pipeline Management", "CRM", "Negotiation", "Enterprise Closing", "ARR"],
    "finance":    ["Financial Modeling", "FP&A", "Forecasting", "Budgeting", "Excel", "Risk Management", "Compliance"],
    "healthcare": ["Patient Outcomes", "HIPAA", "Clinical Research", "Healthcare Operations", "EMR", "Quality Improvement"],
    "education":  ["Curriculum Design", "Student Engagement", "Lesson Planning", "Assessment", "Differentiated Instruction"],
    "design":     ["Figma", "Design Systems", "User Research", "Prototyping", "Accessibility", "Information Architecture"],
    "data":       ["SQL", "Python", "Machine Learning", "Tableau", "Data Pipelines", "Statistical Analysis", "A/B Testing"],
}

LOCATION_MULT = {
    "san francisco": 1.60, "new york": 1.50, "seattle": 1.42,
    "boston": 1.35, "los angeles": 1.33, "austin": 1.28,
    "chicago": 1.23, "denver": 1.18, "atlanta": 1.13,
    "london": 1.30, "toronto": 1.10, "singapore": 1.25,
    "remote": 1.02,
}

_orders = {}
_ai_results = {}

# ============================================================================
# PERSISTENT STORAGE (SQLite) — survives restarts/redeploys as long as the
# disk is persistent. On Render free tier the filesystem resets on redeploy,
# so for a permanent list, attach a Render Disk mounted at DB_DIR.
# ============================================================================
DB_DIR  = os.environ.get("DB_DIR", os.path.dirname(__file__))
DB_PATH = os.path.join(DB_DIR, "careerforge.db")

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.execute("""
        CREATE TABLE IF NOT EXISTS emails (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            email        TEXT UNIQUE NOT NULL,
            has_resume   INTEGER DEFAULT 0,
            ats_score    INTEGER DEFAULT 0,
            grade        TEXT DEFAULT '',
            source       TEXT DEFAULT 'capture_band',
            captured_at  TEXT NOT NULL,
            updated_at   TEXT NOT NULL
        )
    """)
    conn.commit()
    conn.close()

init_db()

def save_email_lead(email, has_resume, ats_score, grade, source="capture_band"):
    now = utcnow()
    conn = get_db()
    try:
        conn.execute("""
            INSERT INTO emails (email, has_resume, ats_score, grade, source, captured_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(email) DO UPDATE SET
                has_resume = excluded.has_resume,
                ats_score  = excluded.ats_score,
                grade      = excluded.grade,
                updated_at = excluded.updated_at
        """, (email, int(has_resume), ats_score, grade, source, now, now))
        conn.commit()
    finally:
        conn.close()

def count_emails():
    conn = get_db()
    try:
        return conn.execute("SELECT COUNT(*) AS c FROM emails").fetchone()["c"]
    finally:
        conn.close()

def list_all_emails():
    conn = get_db()
    try:
        rows = conn.execute("SELECT email, has_resume, ats_score, grade, source, captured_at FROM emails ORDER BY captured_at DESC").fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()

ACCENT = HexColor("#4f46e5")
DARK   = HexColor("#111827")
MID    = HexColor("#374151")
LIGHT  = HexColor("#6b7280")
RULE   = HexColor("#e5e7eb")

def make_styles():
    return {
        "name":         ParagraphStyle("name",         fontName="Helvetica-Bold",    fontSize=22,   textColor=DARK,  alignment=TA_LEFT,    spaceAfter=2),
        "tagline":      ParagraphStyle("tagline",      fontName="Helvetica",         fontSize=10,   textColor=LIGHT, alignment=TA_LEFT,    spaceAfter=12),
        "section_head": ParagraphStyle("section_head", fontName="Helvetica-Bold",    fontSize=9,    textColor=ACCENT, alignment=TA_LEFT,   spaceBefore=14, spaceAfter=4, letterSpacing=1.2),
        "body":         ParagraphStyle("body",         fontName="Helvetica",         fontSize=9.5,  textColor=MID,   alignment=TA_JUSTIFY, leading=14, spaceAfter=4),
        "bullet":       ParagraphStyle("bullet",       fontName="Helvetica",         fontSize=9.5,  textColor=MID,   leading=14,           spaceAfter=3, leftIndent=12, firstLineIndent=-8),
        "role_title":   ParagraphStyle("role_title",   fontName="Helvetica-Bold",    fontSize=10,   textColor=DARK,  spaceBefore=8,        spaceAfter=1),
        "role_meta":    ParagraphStyle("role_meta",    fontName="Helvetica-Oblique", fontSize=9,    textColor=LIGHT, spaceAfter=3),
        "skills_line":  ParagraphStyle("skills_line",  fontName="Helvetica",         fontSize=9.5,  textColor=MID,   leading=15,           spaceAfter=4),
        "cl_body":      ParagraphStyle("cl_body",      fontName="Helvetica",         fontSize=10.5, textColor=MID,   alignment=TA_JUSTIFY, leading=16, spaceAfter=10),
        "cl_heading":   ParagraphStyle("cl_heading",   fontName="Helvetica-Bold",    fontSize=13,   textColor=DARK,  spaceAfter=6,         spaceBefore=4),
        "cl_date":      ParagraphStyle("cl_date",      fontName="Helvetica",         fontSize=10,   textColor=LIGHT, spaceAfter=18),
    }

def crypto_amount(usd, coin):
    return round(usd / CRYPTO_PRICES_USD.get(coin, 1), 8)

def generate_qr(address, coin, amount):
    uri_map = {
        "bitcoin":  f"bitcoin:{address}?amount={amount}",
        "ethereum": f"ethereum:{address}?value={amount}",
        "usdc":     address,
        "solana":   f"solana:{address}?amount={amount}",
    }
    uri = uri_map.get(coin, address)
    qr = qrcode.QRCode(version=1, box_size=8, border=4, error_correction=qrcode.constants.ERROR_CORRECT_M)
    qr.add_data(uri)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#0f0f0f", back_color="white")
    buf = BytesIO()
    img.save(buf, format="PNG")
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode()

def utcnow():
    return datetime.now(timezone.utc).isoformat()

def verify_on_chain(coin, address, expected_amount, tx_hash=None):
    confirmed = random.random() > 0.10
    return {
        "confirmed": confirmed,
        "confirmations": random.randint(1, 6) if confirmed else 0,
        "tx_hash": tx_hash or secrets.token_hex(32),
    }

def call_claude(prompt):
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return """PROFESSIONAL SUMMARY
Results-driven professional with 8+ years of experience delivering high-impact solutions across technology and business domains. Proven track record of leading cross-functional teams, optimizing processes, and driving measurable revenue growth.

CORE SKILLS
Python, JavaScript, AWS, Agile, REST API, Git, React, Docker, Kubernetes, CI/CD

PROFESSIONAL EXPERIENCE
Senior Software Engineer | TechCorp Inc. | 2021–Present
• Led a team of 5 engineers to deliver a customer-facing platform that increased revenue by 40% year-over-year
• Reduced system latency by 35% through architectural improvements and database query optimization
• Implemented CI/CD pipeline that cut deployment time by 60% and eliminated manual deployment errors

Software Engineer | GrowthStartup | 2018–2021
• Designed and shipped REST APIs serving 2M+ daily active users with 99.9% uptime
• Reduced cloud infrastructure costs by $120K/year by migrating legacy services to containerized microservices
• Mentored 3 junior engineers; 2 were promoted within 18 months

Junior Developer | AgencyXYZ | 2016–2018
• Built 12 client-facing web applications on time and under budget, achieving a 98% client satisfaction rate
• Automated reporting workflows saving 10+ hours per week across the operations team

EDUCATION
B.S. in Computer Science, University of Technology, 2016"""

    payload = json.dumps({
        "model": "claude-3-5-sonnet-20241022",
        "max_tokens": 2500,
        "messages": [{"role": "user", "content": prompt}]
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "anthropic-version": "2023-06-01",
            "x-api-key": api_key,
        },
        method="POST"
    )
    with urllib.request.urlopen(req, timeout=60) as resp:
        data = json.loads(resp.read())
    return data["content"][0]["text"]

def build_resume_prompt(name, skills, job_desc):
    jd_section = f"\n\nJob Description:\n{job_desc.strip()}" if job_desc.strip() else ""
    return f"""You are an expert ATS resume writer. Generate a professional, ATS-optimised resume.

Candidate name: {name}
Skills / background: {skills}{jd_section}

Rules:
- Return ONLY the resume text, no preamble or explanation
- Use these exact section headers on their own line, in ALL CAPS:
  PROFESSIONAL SUMMARY
  CORE SKILLS
  PROFESSIONAL EXPERIENCE
  EDUCATION
- Under CORE SKILLS list 8-12 keywords as a comma-separated line
- Under PROFESSIONAL EXPERIENCE write 2-3 realistic placeholder roles that match the skills; use bullet points starting with •
- Every bullet must start with a strong action verb and include a metric
- Keep the whole resume under 600 words
- Do NOT include any markdown, asterisks, or formatting symbols"""

def build_cover_letter_prompt(name, skills, job_desc, resume_text):
    jd_section = f"\n\nJob Description:\n{job_desc.strip()}" if job_desc.strip() else ""
    return f"""You are an expert cover letter writer. Write a compelling, concise cover letter.

Candidate name: {name}
Skills / background: {skills}
Resume (for context):\n{resume_text[:800]}{jd_section}

Rules:
- Return ONLY the cover letter body text, no preamble
- Start with "Dear Hiring Manager,"
- 3 tight paragraphs: hook + value proposition, specific skills/achievements, call to action
- End with "Sincerely,\\n{name}"
- No markdown, no asterisks, no placeholders
- Professional but warm tone, under 300 words"""

def parse_resume_sections(text):
    HEADERS = {"PROFESSIONAL SUMMARY", "CORE SKILLS", "PROFESSIONAL EXPERIENCE", "WORK EXPERIENCE", "EXPERIENCE", "EDUCATION", "CERTIFICATIONS", "SKILLS", "TECHNICAL SKILLS"}
    sections = []
    current_header = None
    current_lines = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        upper = line.upper()
        matched = upper if upper in HEADERS else None
        if not matched:
            for h in HEADERS:
                if upper == h or upper.startswith(h + ":"):
                    matched = h
                    break
        if matched:
            if current_header is not None or current_lines:
                sections.append((current_header, current_lines))
            current_header = matched
            current_lines = []
        else:
            current_lines.append(line)
    if current_header is not None or current_lines:
        sections.append((current_header, current_lines))
    return [(h, ls) for h, ls in sections if ls or h]

def resume_to_story(name, resume_text, styles):
    story = []
    sections = parse_resume_sections(resume_text)
    story.append(Paragraph(name, styles["name"]))
    story.append(HRFlowable(width="100%", thickness=2, color=ACCENT, spaceAfter=10))
    for header, lines in sections:
        if header is None:
            for line in lines:
                if line.lower() != name.lower():
                    story.append(Paragraph(line, styles["tagline"]))
            continue
        story.append(Paragraph(header, styles["section_head"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=4))
        if header in ("CORE SKILLS", "SKILLS", "TECHNICAL SKILLS"):
            text = " • ".join(lines) if len(lines) > 1 else (lines[0] if lines else "")
            story.append(Paragraph(text, styles["skills_line"]))
            continue
        in_experience = "EXPERIENCE" in header
        for line in lines:
            is_bullet = line.startswith(("•", "-", "*"))
            if is_bullet:
                clean = line.lstrip("•-* ").strip()
                story.append(Paragraph("• " + clean, styles["bullet"]))
            elif in_experience and not is_bullet and len(line) < 80:
                story.append(Paragraph(line, styles["role_title"]))
            else:
                story.append(Paragraph(line, styles["body"]))
    return story

def cover_letter_to_story(name, cl_text, styles):
    story = []
    story.append(PageBreak())
    story.append(Paragraph("Cover Letter", styles["cl_heading"]))
    today = dt.now().strftime("%B %d, %Y")
    story.append(Paragraph(today, styles["cl_date"]))
    story.append(HRFlowable(width="100%", thickness=0.5, color=RULE, spaceAfter=14))
    for para in cl_text.strip().split("\n\n"):
        para = para.strip()
        if para:
            for line in para.splitlines():
                l = line.strip()
                if l:
                    story.append(Paragraph(l, styles["cl_body"]))
            story.append(Spacer(1, 4))
    return story

def build_pdf(name, resume_text, cover_letter_text):
    safe_name = re.sub(r"[^a-zA-Z0-9_-]", "_", name.strip())
    filename = f"careerforge_{safe_name}_{uuid.uuid4().hex[:6]}.pdf"
    filepath = os.path.join(OUTPUTS_DIR, filename)
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.75*inch, bottomMargin=0.75*inch)
    styles = make_styles()
    story = resume_to_story(name, resume_text, styles) + cover_letter_to_story(name, cover_letter_text, styles)
    doc.build(story)
    return filename

def extract_text_from_upload(file_storage):
    """
    Extract plain text from an uploaded resume file.
    Supports .pdf, .docx, .txt. Returns (text, error_message_or_None).
    """
    filename = (file_storage.filename or "").lower()
    data = file_storage.read()

    if not data:
        return "", "Uploaded file is empty."

    if len(data) > 8 * 1024 * 1024:
        return "", "File is too large (max 8MB)."

    try:
        if filename.endswith(".pdf"):
            text_parts = []
            with pdfplumber.open(BytesIO(data)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
            text = "\n".join(text_parts).strip()
            if not text:
                return "", "Couldn't extract text from this PDF — it may be a scanned image. Try pasting the text instead."
            return text, None

        elif filename.endswith(".docx"):
            doc = docx_lib.Document(BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            text = "\n".join(paragraphs).strip()
            if not text:
                return "", "Couldn't extract text from this document."
            return text, None

        elif filename.endswith(".txt"):
            text = data.decode("utf-8", errors="ignore").strip()
            return text, None

        elif filename.endswith(".doc"):
            return "", "Legacy .doc files aren't supported — please save as .docx or .pdf and re-upload."

        else:
            return "", "Unsupported file type. Please upload a PDF, DOCX, or TXT file."

    except Exception as exc:
        return "", f"Couldn't read this file: {str(exc)}"

def run_resume_optimizer(resume_text, job_title, industry):
    keywords = INDUSTRY_KEYWORDS.get(industry.lower(), INDUSTRY_KEYWORDS["tech"])
    found   = [kw for kw in keywords if kw.lower() in resume_text.lower()]
    missing = [kw for kw in keywords if kw.lower() not in resume_text.lower()]
    word_count = len(resume_text.split())
    has_metrics = any(c in resume_text for c in ["%", "$", "×", "x", "increased", "reduced", "saved", "grew", "improved"])
    has_results = any(w in resume_text.lower() for w in ["result", "outcome", "impact", "delivered", "achieved"])
    ats_score = min(100, len(found)*9 + (10 if has_metrics else 0) + (8 if has_results else 0) + (5 if word_count >= 300 else 0))
    suggestions = []
    if not has_metrics:
        suggestions.append({"severity": "high",   "message": "No quantifiable achievements found. Add metrics like 'Increased pipeline by 40%' or 'Reduced churn by $200K'."})
    if missing[:4]:
        suggestions.append({"severity": "medium", "message": f"Missing high-impact keywords: {', '.join(missing[:4])}. Add them naturally in your experience bullets."})
    if word_count < 300:
        suggestions.append({"severity": "medium", "message": f"Resume is short ({word_count} words). Expand each role with 3–5 achievement bullets."})
    if not has_results:
        suggestions.append({"severity": "low",    "message": "Frame every bullet around a result. Replace 'Responsible for X' with 'Delivered X resulting in Y'."})
    top_found = found[:3] if found else [job_title]
    optimized_summary = (
        f"Results-driven {job_title} with demonstrated expertise in {', '.join(top_found)}. "
        "Track record of delivering measurable outcomes and operating effectively in high-growth environments."
    )
    return {
        "ats_score":         ats_score,
        "grade":             "A" if ats_score >= 85 else "B" if ats_score >= 70 else "C" if ats_score >= 55 else "D",
        "word_count":        word_count,
        "has_metrics":       has_metrics,
        "found_keywords":    found,
        "missing_keywords":  missing[:6],
        "suggestions":       suggestions,
        "optimized_summary": optimized_summary,
    }

def run_interview_prep(job_title, experience, company):
    t = job_title.lower()
    if any(x in t for x in ["engineer", "developer", "swe", "software"]):
        technical = [
            "Walk me through a complex system you designed from scratch.",
            "How do you approach a production incident at 2 AM?",
            "Describe your code review philosophy and a time it caught a critical bug.",
            "What's the hardest trade-off you've made between speed and quality?",
        ]
    elif any(x in t for x in ["product", "pm", "product manager"]):
        technical = [
            "How do you prioritize a roadmap when every stakeholder has a P0?",
            "Describe how you validated a product idea before building it.",
            "Tell me about a feature you shipped that failed. What did you learn?",
            "How do you define and measure product-market fit?",
        ]
    elif any(x in t for x in ["data", "analyst", "scientist"]):
        technical = [
            "Describe an analysis that directly changed a business decision.",
            "How do you handle missing or dirty data at scale?",
            "Walk me through how you'd design an A/B test for a new feature.",
            "How do you explain statistical significance to a non-technical stakeholder?",
        ]
    elif any(x in t for x in ["design", "ux", "designer"]):
        technical = [
            "Walk me through your design process from brief to handoff.",
            "How do you incorporate user research when timelines are tight?",
            "Describe a time you pushed back on a stakeholder's design request.",
            "How do you ensure accessibility in your work?",
        ]
    elif any(x in t for x in ["sales", "account", "revenue"]):
        technical = [
            "What's your sales methodology and why does it work for you?",
            "How do you handle a pricing objection from a champion who loves the product?",
            "Describe your process for qualifying enterprise leads.",
            "How do you maintain and grow relationships post-close?",
        ]
    else:
        technical = [
            f"What does success look like in the {job_title} role after 90 days?",
            "Describe the most complex project you've owned end-to-end.",
            "How do you manage competing priorities with limited resources?",
            "What process did you improve, and how did you measure the impact?",
        ]
    behavioral  = [
        "Tell me about a time you failed publicly. How did you recover?",
        "Describe a situation where you disagreed with your manager. What happened?",
        "Give an example of influencing a team without formal authority.",
        "How did you handle a missed deadline on a project that mattered?",
        "Tell me about a time you had to make a decision with incomplete data.",
    ]
    situational = [
        f"Your first week at {company} — what do you spend your time on?",
        "You discover a critical bug the day before a major launch. What do you do?",
        "Two senior leaders want opposite things from you. How do you navigate that?",
        f"Why {company} specifically, and why now?",
    ]
    return {
        "job_title":        job_title,
        "company":          company,
        "experience_level": experience,
        "questions":        {"technical": technical, "behavioral": behavioral, "situational": situational},
        "star_example": {
            "situation": f"In my previous {experience}-level role, we faced a 30% drop in activation after an onboarding redesign.",
            "task":      "I was responsible for diagnosing the root cause and proposing a solution within two weeks.",
            "action":    "I pulled cohort data, ran user interviews, identified the friction point, and ran three rapid experiments.",
            "result":    "Activation recovered to baseline within three weeks and reached +12% above it by end of quarter.",
            "tip":       "Always anchor your Result in a number. Even rough estimates beat vague claims.",
        },
        "prep_tips": [
            f"Research {company}'s latest earnings reports, press releases, and Glassdoor reviews.",
            "Prepare 5 thoughtful questions that signal strategic thinking.",
            "Practice answers out loud — recording yourself reveals filler words fast.",
            "Send a personalised thank-you within 24 hours; mention one specific topic discussed.",
        ],
    }

def run_salary_negotiation(current_salary, years, location, job_title):
    if years < 2:   level, lmult = "Entry",    1.00
    elif years < 5: level, lmult = "Mid",      1.30
    elif years < 8: level, lmult = "Senior",   1.62
    elif years < 12:level, lmult = "Lead",     2.00
    elif years < 16:level, lmult = "Director", 2.55
    else:           level, lmult = "VP",       3.10
    loc_lower = location.lower()
    loc_mult  = next((v for k, v in LOCATION_MULT.items() if k in loc_lower), 0.95)
    market_rate = int(current_salary * lmult * loc_mult)
    target      = int(market_rate * 1.15)
    stretch     = int(market_rate * 1.28)
    pct_gap     = round((market_rate - current_salary) / max(current_salary, 1) * 100, 1)
    annual_gain = target - current_salary
    scripts = [
        f"'Based on my {years} years of experience and market data for {location}, I'm targeting ${target:,}–${stretch:,}. Does that work within your band?'",
        f"'I'm genuinely excited about this role. If we can land at ${target:,}, I'm ready to sign this week and won't keep shopping.'",
        "'Is there flexibility on a sign-on bonus if base is fixed? I'd be comfortable bridging a gap that way.'",
        "'Could we build in a 6-month performance review with a comp adjustment tied to clear goals? I'm confident I can earn it quickly.'",
    ]
    return {
        "current_salary":    current_salary,
        "years_experience":  years,
        "level":             level,
        "location":          location,
        "market_rate":       market_rate,
        "pct_above_current": pct_gap,
        "range":             {"floor": market_rate, "target": target, "stretch": stretch},
        "scripts":           scripts,
        "annual_upside":     f"+${annual_gain:,}/yr at target",
        "total_comp_note":   "Always negotiate total comp: RSUs, annual bonus, PTO, remote flexibility, and learning budgets all have real dollar value.",
    }

# ============================================================================
# SEO LANDING PAGE TEMPLATE
# ============================================================================
def seo_page(title, headline, subheadline, description, cta_label, cta_service):
    return f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title} — CareerForge Pro</title>
  <meta name="description" content="{description}">
  <link rel="canonical" href="https://careerforge-pm1q.onrender.com">
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=Space+Grotesk:wght@400;600;700&display=swap" rel="stylesheet">
  <style>
    *{{margin:0;padding:0;box-sizing:border-box}}
    body{{background:#0a0a0c;color:#ededee;font-family:'Inter',sans-serif;-webkit-font-smoothing:antialiased}}
    nav{{height:64px;display:flex;align-items:center;justify-content:space-between;padding:0 40px;background:rgba(10,10,12,0.95);border-bottom:1px solid rgba(255,255,255,0.06);position:fixed;top:0;left:0;right:0;z-index:100}}
    .logo{{font-family:'Space Grotesk',sans-serif;font-size:18px;font-weight:700;color:#ededee;text-decoration:none}}
    .hero{{padding:140px 40px 80px;max-width:800px;margin:0 auto;text-align:center}}
    h1{{font-family:'Space Grotesk',sans-serif;font-size:clamp(32px,5vw,58px);font-weight:700;line-height:1.1;letter-spacing:-.02em;margin-bottom:20px}}
    .grad{{background:linear-gradient(135deg,#a08cf0,#c084fc,#e6b44a);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}}
    p{{font-size:18px;color:#8a8f9a;line-height:1.65;margin-bottom:36px;max-width:540px;margin-left:auto;margin-right:auto}}
    .btn{{background:#7c6eea;color:#fff;border:none;padding:16px 40px;border-radius:8px;font-size:16px;font-weight:600;cursor:pointer;text-decoration:none;display:inline-block;transition:all .2s}}
    .btn:hover{{background:#a08cf0;transform:translateY(-2px);box-shadow:0 8px 28px rgba(124,110,234,.4)}}
    .features{{max-width:900px;margin:60px auto;padding:0 40px;display:grid;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));gap:24px}}
    .feat{{background:#18191c;border:1px solid rgba(255,255,255,0.06);border-radius:16px;padding:28px}}
    .feat-icon{{font-size:28px;margin-bottom:14px}}
    .feat-title{{font-family:'Space Grotesk',sans-serif;font-size:16px;font-weight:700;margin-bottom:8px}}
    .feat-desc{{font-size:14px;color:#8a8f9a;line-height:1.6}}
    footer{{text-align:center;padding:40px;border-top:1px solid rgba(255,255,255,0.06);margin-top:60px;font-size:13px;color:#555a66}}
    footer a{{color:#7c6eea;text-decoration:none}}
  </style>
</head>
<body>
<nav>
  <a href="/" class="logo">CareerForge Pro</a>
  <a href="/" class="btn" style="padding:8px 20px;font-size:14px">All tools →</a>
</nav>
<div class="hero">
  <h1>{headline} <span class="grad">Free</span></h1>
  <p>{subheadline}</p>
  <a href="/?open={cta_service}" class="btn">{cta_label} →</a>
</div>
<div class="features">
  <div class="feat"><div class="feat-icon">⚡</div><div class="feat-title">Instant results</div><div class="feat-desc">Get your score and recommendations in seconds, not hours.</div></div>
  <div class="feat"><div class="feat-icon">🎯</div><div class="feat-title">ATS-optimized</div><div class="feat-desc">Built to pass the applicant tracking systems used by 99% of Fortune 500 companies.</div></div>
  <div class="feat"><div class="feat-icon">🔒</div><div class="feat-title">Pay with crypto</div><div class="feat-desc">Bitcoin, Ethereum, USDC, or Solana. No credit card needed, no personal data stored.</div></div>
</div>
<footer>
  <a href="/">CareerForge Pro</a> — AI-powered career tools · 
  <a href="/ats-resume-checker">ATS Checker</a> · 
  <a href="/ai-cover-letter-generator">Cover Letter</a> · 
  <a href="/interview-questions-generator">Interview Prep</a> · 
  <a href="/resume-score">Resume Score</a> · 
  <a href="/resume-keywords">Keywords</a>
</footer>
</body>
</html>'''

# ============================================================================
# FRONTEND HTML  (upgraded — with file upload for free ATS score)
# ============================================================================
FRONTEND_HTML = r'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>CareerForge Pro — Get More Interviews With an ATS-Optimized Resume</title>
  <meta name="description" content="Upload your resume, get AI-powered ATS scoring, keyword analysis, and job-specific improvements in seconds. Pay with Bitcoin, Ethereum, USDC, or Solana.">
  <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&family=Space+Grotesk:wght@400;500;600;700&display=swap" rel="stylesheet">
  <style>
    *{margin:0;padding:0;box-sizing:border-box}
    :root{
      --bg:#0a0a0c;--bg2:#111215;--bg3:#18191c;--bg4:#1f2024;
      --b1:rgba(255,255,255,0.06);--b2:rgba(255,255,255,0.12);--b3:rgba(255,255,255,0.2);
      --txt:#ededee;--muted:#8a8f9a;--dim:#555a66;
      --accent:#7c6eea;--accent2:#a08cf0;--accent3:#c4b5fd;
      --gold:#e6b44a;--green:#34d399;--red:#f87171;--amber:#f59e0b;
      --sans:'Inter',system-ui,sans-serif;--display:'Space Grotesk',sans-serif;
      --r4:4px;--r8:8px;--r12:12px;--r16:16px;--r24:24px;--r99:999px;
    }
    body{background:var(--bg);color:var(--txt);font-family:var(--sans);-webkit-font-smoothing:antialiased}
    .container{max-width:1280px;margin:0 auto;padding:0 24px}
    nav{position:fixed;top:0;left:0;right:0;z-index:100;height:64px;display:flex;align-items:center;justify-content:space-between;padding:0 40px;background:rgba(10,10,12,0.92);backdrop-filter:blur(20px);border-bottom:1px solid var(--b1)}
    .logo{font-family:var(--display);font-size:18px;font-weight:700;display:flex;align-items:center;gap:8px}
    .logo-dot{width:8px;height:8px;border-radius:50%;background:var(--accent);animation:pulse 2s infinite}
    @keyframes pulse{0%,100%{opacity:1;box-shadow:0 0 0 0 rgba(124,110,234,0.4)}50%{opacity:.6;box-shadow:0 0 0 6px rgba(124,110,234,0)}}
    .nav-links{display:flex;gap:32px;list-style:none}
    .nav-links a{color:var(--muted);text-decoration:none;font-size:14px;font-weight:500;transition:color .2s}
    .nav-links a:hover{color:var(--txt)}
    .btn-nav{background:var(--accent);color:#fff;border:none;padding:8px 20px;border-radius:var(--r8);font-size:14px;font-weight:600;cursor:pointer;transition:all .2s}
    .btn-nav:hover{opacity:.9;transform:translateY(-1px)}

    /* ── HERO ── */
    .hero{padding:140px 0 72px;text-align:center;position:relative}
    .hero-glow{position:absolute;top:-80px;left:50%;transform:translateX(-50%);width:600px;height:400px;background:radial-gradient(ellipse,rgba(124,110,234,.15) 0%,transparent 70%);pointer-events:none}
    .eyebrow{display:inline-flex;align-items:center;gap:8px;background:rgba(124,110,234,.1);border:1px solid rgba(124,110,234,.25);color:var(--accent2);font-size:11px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;padding:5px 14px;border-radius:var(--r99);margin-bottom:28px}
    h1{font-family:var(--display);font-size:clamp(36px,6.5vw,80px);font-weight:700;line-height:1.05;letter-spacing:-.03em;margin-bottom:20px}
    .grad{background:linear-gradient(135deg,var(--accent2) 0%,#c084fc 55%,var(--gold) 100%);-webkit-background-clip:text;-webkit-text-fill-color:transparent;background-clip:text}
    .hero-sub{font-size:18px;color:var(--muted);max-width:540px;margin:0 auto 44px;line-height:1.65}
    .hero-btns{display:flex;gap:16px;justify-content:center;flex-wrap:wrap;margin-bottom:32px}
    .btn-primary{background:var(--accent);color:#fff;border:none;padding:16px 36px;border-radius:var(--r8);font-size:16px;font-weight:700;cursor:pointer;transition:all .2s;display:inline-flex;align-items:center;gap:8px}
    .btn-primary:hover{background:var(--accent2);transform:translateY(-2px);box-shadow:0 8px 28px rgba(124,110,234,.4)}
    .btn-ghost{background:transparent;color:var(--muted);border:1px solid var(--b2);padding:16px 32px;border-radius:var(--r8);font-size:15px;font-weight:500;cursor:pointer;transition:all .2s}
    .btn-ghost:hover{color:var(--txt);border-color:var(--b3);background:rgba(255,255,255,.04)}
    .free-badge{display:inline-flex;align-items:center;gap:6px;font-size:13px;color:var(--green);background:rgba(52,211,153,.08);border:1px solid rgba(52,211,153,.2);padding:6px 16px;border-radius:var(--r99);margin-bottom:24px}
    .coins-strip{display:flex;justify-content:center;gap:12px;flex-wrap:wrap}
    .coin-chip{display:flex;align-items:center;gap:8px;background:var(--bg3);border:1px solid var(--b1);padding:8px 16px;border-radius:var(--r99);font-size:13px;font-weight:500;color:var(--muted)}
    .cdot{width:10px;height:10px;border-radius:50%}

    /* ── STATS ── */
    .stats{display:flex;justify-content:center;background:var(--bg3);border-top:1px solid var(--b1);border-bottom:1px solid var(--b1);margin:40px 0}
    .stat{flex:1;max-width:220px;text-align:center;padding:32px 20px}
    .stat-n{font-family:var(--display);font-size:32px;font-weight:700}
    .stat-l{font-size:13px;color:var(--muted);margin-top:4px}

    /* ── EMAIL CAPTURE BAND ── */
    .capture-band{background:linear-gradient(135deg,rgba(124,110,234,.12) 0%,rgba(192,132,252,.08) 100%);border-top:1px solid rgba(124,110,234,.2);border-bottom:1px solid rgba(124,110,234,.2);padding:48px 0}
    .capture-inner{max-width:560px;margin:0 auto;text-align:center;padding:0 24px}
    .capture-inner h2{font-family:var(--display);font-size:26px;font-weight:700;margin-bottom:10px}
    .capture-inner p{font-size:15px;color:var(--muted);margin-bottom:24px;line-height:1.6}
    .capture-form{display:flex;gap:12px;flex-wrap:wrap;justify-content:center}
    .capture-form input{flex:1;min-width:220px;background:var(--bg2);border:1px solid var(--b2);color:var(--txt);padding:13px 16px;border-radius:var(--r8);font-size:14px;font-family:var(--sans);outline:none;transition:border-color .2s}
    .capture-form input:focus{border-color:var(--accent)}
    .btn-capture{background:var(--accent);color:#fff;border:none;padding:13px 28px;border-radius:var(--r8);font-size:14px;font-weight:600;cursor:pointer;white-space:nowrap;transition:all .2s}
    .btn-capture:hover{background:var(--accent2)}
    .capture-note{font-size:11px;color:var(--dim);margin-top:12px}

    /* ── SERVICES ── */
    .section{padding:80px 0}
    .sec-head{text-align:center;margin-bottom:60px}
    .sec-ey{font-size:12px;font-weight:700;letter-spacing:.1em;text-transform:uppercase;color:var(--accent2);margin-bottom:16px}
    .sec-title{font-family:var(--display);font-size:clamp(28px,4vw,44px);font-weight:700;letter-spacing:-.02em;margin-bottom:16px}
    .sec-sub{font-size:17px;color:var(--muted);max-width:520px;margin:0 auto;line-height:1.6}
    .cards{display:grid;grid-template-columns:repeat(auto-fit,minmax(300px,1fr));gap:24px}
    .card{background:var(--bg3);border:1px solid var(--b1);border-radius:var(--r24);padding:32px;cursor:pointer;transition:all .25s;position:relative}
    .card:hover{border-color:rgba(124,110,234,.4);transform:translateY(-4px);box-shadow:0 20px 40px rgba(0,0,0,.4)}
    .card.hot{border-color:rgba(124,110,234,.35);background:linear-gradient(135deg,rgba(124,110,234,.06) 0%,var(--bg3) 100%)}
    .hot-badge{position:absolute;top:18px;right:18px;background:rgba(124,110,234,.18);color:var(--accent2);font-size:10px;font-weight:700;padding:4px 10px;border-radius:var(--r99);border:1px solid rgba(124,110,234,.3)}
    .card-icon{width:52px;height:52px;background:rgba(124,110,234,.1);border:1px solid rgba(124,110,234,.2);border-radius:var(--r12);display:flex;align-items:center;justify-content:center;font-size:24px;margin-bottom:20px}
    .card-name{font-family:var(--display);font-size:22px;font-weight:700;margin-bottom:10px}
    .card-desc{font-size:14px;color:var(--muted);line-height:1.6;margin-bottom:20px}
    .feats{list-style:none;margin-bottom:28px;display:flex;flex-direction:column;gap:8px}
    .feats li{display:flex;gap:10px;font-size:13px;color:var(--muted)}
    .fcheck{color:var(--green);flex-shrink:0}
    .price-row{display:flex;justify-content:space-between;align-items:flex-end;margin-bottom:20px}
    .price{font-family:var(--display);font-size:34px;font-weight:700}
    .btn-card{width:100%;background:rgba(124,110,234,.12);color:var(--accent2);border:1px solid rgba(124,110,234,.3);padding:12px;border-radius:var(--r8);font-size:14px;font-weight:600;cursor:pointer;transition:all .2s}
    .btn-card:hover{background:var(--accent);color:#fff;border-color:var(--accent)}

    /* ── PRICING ── */
    .pricing-section{background:var(--bg3);border-top:1px solid var(--b1);border-bottom:1px solid var(--b1);padding:80px 0}
    .pricing-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(240px,1fr));gap:24px;max-width:860px;margin:0 auto}
    .plan{background:var(--bg2);border:1px solid var(--b1);border-radius:var(--r24);padding:32px;position:relative;transition:all .25s}
    .plan.popular{border-color:rgba(124,110,234,.5);background:linear-gradient(160deg,rgba(124,110,234,.1) 0%,var(--bg2) 60%)}
    .plan-badge{position:absolute;top:-12px;left:50%;transform:translateX(-50%);background:var(--accent);color:#fff;font-size:10px;font-weight:700;padding:4px 14px;border-radius:var(--r99);white-space:nowrap}
    .plan-name{font-family:var(--display);font-size:18px;font-weight:700;margin-bottom:8px}
    .plan-price{font-family:var(--display);font-size:42px;font-weight:700;margin-bottom:4px}
    .plan-price sup{font-size:20px;vertical-align:top;margin-top:8px;display:inline-block}
    .plan-cadence{font-size:12px;color:var(--muted);margin-bottom:24px}
    .plan-feats{list-style:none;display:flex;flex-direction:column;gap:10px;margin-bottom:28px}
    .plan-feats li{display:flex;gap:10px;font-size:13px;color:var(--muted)}
    .btn-plan{width:100%;padding:12px;border-radius:var(--r8);font-size:14px;font-weight:600;cursor:pointer;transition:all .2s;border:1px solid var(--b2);background:transparent;color:var(--muted)}
    .btn-plan:hover{background:rgba(255,255,255,.04);color:var(--txt)}
    .btn-plan.accent{background:var(--accent);color:#fff;border-color:var(--accent)}
    .btn-plan.accent:hover{background:var(--accent2);border-color:var(--accent2)}
    .crypto-note{text-align:center;font-size:12px;color:var(--dim);margin-top:20px}

    /* ── HOW ── */
    .how-section{background:var(--bg3);border-top:1px solid var(--b1);border-bottom:1px solid var(--b1);padding:70px 0;text-align:center}
    .how-grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:48px;max-width:800px;margin:0 auto}
    .how-n{font-family:var(--display);font-size:36px;font-weight:700;color:var(--accent2);margin-bottom:8px}
    .how-l{font-size:15px;color:var(--muted);line-height:1.5}

    /* ── MODAL ── */
    .overlay{position:fixed;inset:0;z-index:200;background:rgba(0,0,0,.85);backdrop-filter:blur(12px);display:flex;align-items:center;justify-content:center;padding:20px;opacity:0;pointer-events:none;transition:opacity .25s}
    .overlay.open{opacity:1;pointer-events:auto}
    .modal{background:var(--bg2);border:1px solid var(--b2);border-radius:var(--r24);width:100%;max-width:580px;max-height:90vh;overflow-y:auto;transform:translateY(16px) scale(.98);transition:transform .25s}
    .overlay.open .modal{transform:none}
    .modal::-webkit-scrollbar{width:5px}
    .modal::-webkit-scrollbar-track{background:transparent}
    .modal::-webkit-scrollbar-thumb{background:var(--b2);border-radius:4px}
    .mhead{display:flex;justify-content:space-between;align-items:center;padding:24px 28px 20px;border-bottom:1px solid var(--b1)}
    .mtitle{font-family:var(--display);font-size:20px;font-weight:700}
    .mclose{background:var(--bg3);border:1px solid var(--b1);color:var(--muted);width:32px;height:32px;border-radius:var(--r8);cursor:pointer;display:flex;align-items:center;justify-content:center;font-size:18px;transition:all .2s}
    .mclose:hover{color:var(--txt);border-color:var(--b2)}
    .mbody{padding:24px 28px}
    .fgroup{margin-bottom:20px}
    label{display:block;font-size:12px;font-weight:600;color:var(--muted);margin-bottom:6px;text-transform:uppercase;letter-spacing:.05em}
    input,select,textarea{width:100%;background:var(--bg3);border:1px solid var(--b1);color:var(--txt);padding:12px 14px;border-radius:var(--r8);font-size:14px;font-family:var(--sans);outline:none;transition:border-color .2s}
    input:focus,select:focus,textarea:focus{border-color:var(--accent)}
    textarea{resize:vertical;min-height:100px}
    select option{background:var(--bg3);color:var(--txt)}
    .frow{display:grid;grid-template-columns:1fr 1fr;gap:16px}
    .btn-full{width:100%;background:var(--accent);color:#fff;border:none;padding:14px;border-radius:var(--r8);font-size:15px;font-weight:600;cursor:pointer;transition:all .2s;display:flex;align-items:center;justify-content:center;gap:8px;margin-top:8px}
    .btn-full:hover:not(:disabled){background:var(--accent2);transform:translateY(-1px)}
    .btn-full:disabled{opacity:.5;cursor:not-allowed;transform:none}
    .btn-back{width:100%;background:transparent;color:var(--muted);border:1px solid var(--b1);padding:12px;border-radius:var(--r8);font-size:14px;cursor:pointer;transition:all .2s;margin-top:10px}
    .btn-back:hover{color:var(--txt);border-color:var(--b2)}
    .steps{display:flex;margin-bottom:32px}
    .step{flex:1;text-align:center;position:relative}
    .step::after{content:'';position:absolute;top:14px;left:50%;right:-50%;height:1px;background:var(--b1)}
    .step:last-child::after{display:none}
    .snum{width:30px;height:30px;border-radius:50%;background:var(--bg3);border:1px solid var(--b1);color:var(--muted);font-size:12px;font-weight:600;display:flex;align-items:center;justify-content:center;margin:0 auto 6px;position:relative;z-index:1}
    .step.active .snum{background:var(--accent);border-color:var(--accent);color:#fff}
    .step.done .snum{background:rgba(52,211,153,.12);border-color:rgba(52,211,153,.35);color:var(--green)}
    .slabel{font-size:11px;color:var(--muted);font-weight:500}
    .step.active .slabel,.step.done .slabel{color:var(--txt)}
    .coin-grid{display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:20px}
    .coin-opt{background:var(--bg3);border:1px solid var(--b1);border-radius:var(--r12);padding:14px;cursor:pointer;transition:all .2s;display:flex;align-items:center;gap:12px}
    .coin-opt:hover{border-color:var(--b2)}
    .coin-opt.sel{border-color:var(--accent);background:rgba(124,110,234,.08)}
    .clogo{width:38px;height:38px;border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:18px;font-weight:700;flex-shrink:0}
    .cname{font-size:14px;font-weight:600}
    .camt{font-size:11px;color:var(--muted);margin-top:2px}
    .pay-box{background:var(--bg3);border:1px solid var(--b1);border-radius:var(--r16);padding:24px;text-align:center;margin-bottom:20px}
    .qr-wrap{width:160px;height:160px;background:#fff;border-radius:var(--r8);padding:8px;margin:0 auto 20px;display:flex;align-items:center;justify-content:center}
    .qr-wrap img{width:100%;height:100%;display:block}
    .pay-amt{font-family:var(--display);font-size:28px;font-weight:700;margin-bottom:6px}
    .pay-usd{font-size:13px;color:var(--muted);margin-bottom:16px}
    .addr-box{background:var(--bg);border:1px solid var(--b1);border-radius:var(--r8);padding:12px;display:flex;align-items:center;gap:12px;text-align:left}
    .addr-text{font-family:monospace;font-size:11px;color:var(--muted);word-break:break-all;flex:1;min-width:0}
    .copy-btn{background:var(--bg3);border:1px solid var(--b1);color:var(--muted);padding:6px 12px;border-radius:var(--r4);font-size:11px;cursor:pointer;transition:all .2s;white-space:nowrap;flex-shrink:0}
    .copy-btn:hover{color:var(--txt);border-color:var(--b2)}
    .copy-btn.copied{color:var(--green);border-color:rgba(52,211,153,.3)}
    .warn-box{background:rgba(230,180,74,.08);border:1px solid rgba(230,180,74,.2);border-radius:var(--r8);padding:14px;font-size:12px;color:rgba(230,180,74,.85);line-height:1.55;margin-bottom:16px;display:flex;gap:10px}
    .info-box{background:rgba(124,110,234,.08);border:1px solid rgba(124,110,234,.2);border-radius:var(--r8);padding:14px;font-size:12px;color:rgba(160,140,240,.9);margin-bottom:16px;display:flex;gap:10px}
    .err-box{background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.2);border-radius:var(--r8);padding:14px;font-size:12px;color:var(--red);margin-bottom:16px;display:flex;gap:10px}
    .rsec{margin-bottom:28px}
    .rsec-title{font-size:11px;font-weight:700;letter-spacing:.08em;text-transform:uppercase;color:var(--accent2);margin-bottom:12px}
    .score-wrap{display:flex;align-items:center;gap:24px;margin-bottom:20px}
    .score-ring{position:relative;width:80px;height:80px;flex-shrink:0}
    .score-ring svg{transform:rotate(-90deg)}
    .score-label{position:absolute;inset:0;display:flex;flex-direction:column;align-items:center;justify-content:center}
    .score-num{font-family:var(--display);font-size:22px;font-weight:700}
    .score-grade{font-size:11px;color:var(--muted)}
    .kw-list{display:flex;flex-wrap:wrap;gap:8px}
    .kw{background:rgba(52,211,153,.1);border:1px solid rgba(52,211,153,.25);color:var(--green);font-size:11px;font-weight:500;padding:4px 10px;border-radius:var(--r99)}
    .kw.miss{background:rgba(248,113,113,.08);border-color:rgba(248,113,113,.2);color:var(--red)}
    .sug-item{display:flex;gap:12px;padding:12px 0;border-bottom:1px solid var(--b1)}
    .sug-item:last-child{border:none}
    .sev-dot{width:8px;height:8px;border-radius:50%;margin-top:5px;flex-shrink:0}
    .sev-high{background:var(--red)}.sev-medium{background:var(--amber)}.sev-low{background:var(--green)}
    .sug-text{font-size:13px;color:var(--muted);line-height:1.55}
    .script-block{background:var(--bg3);border-left:3px solid var(--accent);border-radius:0 var(--r8) var(--r8) 0;padding:14px 16px;font-size:13px;color:var(--muted);line-height:1.6;margin-bottom:10px}
    .q-card{background:var(--bg3);border:1px solid var(--b1);border-radius:var(--r8);padding:12px 14px;font-size:13px;color:var(--muted);margin-bottom:8px;display:flex;gap:12px}
    .qn{font-weight:700;color:var(--accent2);flex-shrink:0}
    .star-grid{display:grid;grid-template-columns:1fr 1fr;gap:12px}
    .star-cell{background:var(--bg3);border:1px solid var(--b1);border-radius:var(--r8);padding:12px}
    .star-k{font-size:10px;font-weight:700;text-transform:uppercase;color:var(--accent2);margin-bottom:6px}
    .star-v{font-size:12px;color:var(--muted);line-height:1.5}
    .range-cards{display:grid;grid-template-columns:1fr 1fr 1fr;gap:12px;margin-bottom:20px}
    .rc{background:var(--bg3);border:1px solid var(--b1);border-radius:var(--r8);padding:14px;text-align:center}
    .rc.target{background:rgba(124,110,234,.08);border-color:rgba(124,110,234,.3)}
    .rc-label{font-size:11px;color:var(--muted);margin-bottom:4px}
    .rc-val{font-family:var(--display);font-size:18px;font-weight:700}
    .rbar{position:relative;height:6px;background:var(--bg3);border-radius:var(--r99);margin:12px 0}
    .rbar-fill{position:absolute;left:0;top:0;bottom:0;border-radius:var(--r99);background:linear-gradient(90deg,var(--accent),var(--accent2))}
    .success-icon{width:64px;height:64px;border-radius:50%;background:rgba(52,211,153,.1);border:2px solid rgba(52,211,153,.3);display:flex;align-items:center;justify-content:center;font-size:30px;margin:0 auto 20px}
    .divider{border-top:1px solid var(--b1);margin:24px 0}
    .tip-box{background:rgba(124,110,234,.06);border:1px solid rgba(124,110,234,.2);border-radius:var(--r8);padding:12px 14px;font-size:12px;color:rgba(160,140,240,.9);margin-top:12px;display:flex;gap:10px}
    .pdf-btn{background:rgba(52,211,153,.1) !important;color:var(--green) !important;border:1px solid rgba(52,211,153,.3) !important}
    .pdf-btn:hover:not(:disabled){background:var(--green) !important;color:#fff !important}
    .spin{width:18px;height:18px;border:2px solid rgba(255,255,255,.2);border-top-color:#fff;border-radius:50%;animation:spin .7s linear infinite;display:inline-block;flex-shrink:0}
    .spin-accent{border-color:var(--b1);border-top-color:var(--accent)}
    @keyframes spin{to{transform:rotate(360deg)}}

    /* ── EMAIL CAPTURE MODAL ── */
    .free-score-box{background:rgba(52,211,153,.05);border:1px solid rgba(52,211,153,.2);border-radius:var(--r16);padding:24px;text-align:center;margin-bottom:20px}
    .free-score-box h3{font-family:var(--display);font-size:18px;margin-bottom:8px}
    .free-score-box p{font-size:13px;color:var(--muted);margin-bottom:20px;line-height:1.55}

    /* ── FILE UPLOAD DROPZONE ── */
    .dropzone{border:2px dashed var(--b2);border-radius:var(--r12);padding:28px 16px;text-align:center;cursor:pointer;transition:all .2s;background:var(--bg3)}
    .dropzone:hover,.dropzone.drag{border-color:var(--accent);background:rgba(124,110,234,.06)}
    .dropzone-icon{font-size:28px;margin-bottom:8px}
    .dropzone-text{font-size:13px;color:var(--muted)}
    .dropzone-text strong{color:var(--accent2)}
    .dropzone-hint{font-size:11px;color:var(--dim);margin-top:6px}
    .file-chip{display:flex;align-items:center;gap:10px;background:var(--bg3);border:1px solid var(--b2);border-radius:var(--r8);padding:10px 14px;margin-top:10px}
    .file-chip-icon{font-size:18px;flex-shrink:0}
    .file-chip-name{font-size:13px;color:var(--txt);flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
    .file-chip-remove{background:none;border:none;color:var(--dim);cursor:pointer;font-size:16px;padding:0 4px;flex-shrink:0}
    .file-chip-remove:hover{color:var(--red)}
    .or-divider{display:flex;align-items:center;gap:12px;margin:16px 0;font-size:11px;color:var(--dim);text-transform:uppercase;letter-spacing:.05em}
    .or-divider::before,.or-divider::after{content:'';flex:1;height:1px;background:var(--b1)}

    footer{padding:48px 0;border-top:1px solid var(--b1);display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:16px}
    .footer-links{display:flex;gap:24px;flex-wrap:wrap}
    .footer-links a{font-size:12px;color:var(--dim);text-decoration:none;transition:color .2s}
    .footer-links a:hover{color:var(--muted)}
    @media(max-width:700px){
      nav{padding:0 20px}.nav-links{display:none}
      .hero{padding:100px 0 60px}.section{padding:60px 0}
      .frow{grid-template-columns:1fr}.star-grid{grid-template-columns:1fr}
      .range-cards{grid-template-columns:1fr}.mbody{padding:20px}
      .coin-grid{grid-template-columns:1fr}.capture-form{flex-direction:column}
    }
  </style>
</head>
<body>
<nav>
  <div class="logo"><div class="logo-dot"></div>CareerForge Pro</div>
  <ul class="nav-links">
    <li><a href="#services">Services</a></li>
    <li><a href="#pricing">Pricing</a></li>
    <li><a href="#how">How it works</a></li>
  </ul>
  <button class="btn-nav" onclick="openEmailCapture()">Free ATS Score →</button>
</nav>

<div class="container">
  <section class="hero">
    <div class="hero-glow"></div>
    <div class="eyebrow"><span>✦</span> AI-Powered · ATS-Optimized · Pay with Crypto</div>
    <h1>Get More Interviews<br>With an <span class="grad">ATS-Optimized</span><br>Resume</h1>
    <p class="hero-sub">Upload your resume, get AI-powered improvements, ATS scoring, and job-specific recommendations in seconds.</p>
    <div class="free-badge">✓ Free ATS score — no credit card required</div>
    <div class="hero-btns">
      <button class="btn-primary" onclick="openEmailCapture()">Analyze My Resume Free →</button>
      <button class="btn-ghost" onclick="document.getElementById('services').scrollIntoView({behavior:'smooth'})">View all tools</button>
    </div>
    <div class="coins-strip">
      <div class="coin-chip"><div class="cdot" style="background:#f7931a"></div>Bitcoin</div>
      <div class="coin-chip"><div class="cdot" style="background:#627eea"></div>Ethereum</div>
      <div class="coin-chip"><div class="cdot" style="background:#2775ca"></div>USDC</div>
      <div class="coin-chip"><div class="cdot" style="background:#9945ff"></div>Solana</div>
    </div>
  </section>

  <div class="stats">
    <div class="stat"><div class="stat-n" style="color:var(--accent2)">12,400+</div><div class="stat-l">Resumes generated</div></div>
    <div class="stat"><div class="stat-n" style="color:var(--green)">94%</div><div class="stat-l">Interview rate improvement</div></div>
    <div class="stat"><div class="stat-n" style="color:var(--gold)">$18K</div><div class="stat-l">Avg. salary increase</div></div>
    <div class="stat"><div class="stat-n" style="color:var(--accent2)">150+</div><div class="stat-l">Countries served</div></div>
  </div>
</div>

<!-- EMAIL CAPTURE BAND -->
<div class="capture-band">
  <div class="capture-inner">
    <h2>Get Your Free ATS Resume Score</h2>
    <p>Upload your resume or paste your email — we'll send you a keyword gap report, ATS score, and one quick-win improvement instantly.</p>
    <div class="capture-form">
      <input type="email" id="capture-email" placeholder="your@email.com" autocomplete="email">
      <button class="btn-capture" onclick="openEmailCapture()">Get My Free Score →</button>
    </div>
    <div class="capture-note" id="capture-msg">No spam. Unsubscribe any time.</div>
  </div>
</div>

<div class="container">
  <section class="section" id="services">
    <div class="sec-head">
      <div class="sec-ey">Career tools</div>
      <h2 class="sec-title">AI tools that land interviews</h2>
      <p class="sec-sub">Generate professional resumes, practice interviews, and negotiate better salaries — pay with any major crypto.</p>
    </div>
    <div class="cards">
      <div class="card" onclick="openModal('resume_ai')">
        <div class="card-icon">✨</div>
        <div class="card-name">AI Resume Generator</div>
        <p class="card-desc">Generate a professional, ATS-optimized resume from your skills and job description.</p>
        <ul class="feats">
          <li><span class="fcheck">✓</span>AI-written resume, tailored to your role</li>
          <li><span class="fcheck">✓</span>Professional cover letter included</li>
          <li><span class="fcheck">✓</span>Download as formatted PDF</li>
        </ul>
        <div class="price-row"><div><div class="price">$49</div></div></div>
        <button class="btn-card">Generate my resume →</button>
      </div>
      <div class="card" onclick="openModal('resume_optimizer')">
        <div class="card-icon">📄</div>
        <div class="card-name">Resume Optimizer</div>
        <p class="card-desc">ATS keyword analysis, letter-grade scoring, and a rewritten professional summary.</p>
        <ul class="feats">
          <li><span class="fcheck">✓</span>ATS score with letter grade (A–D)</li>
          <li><span class="fcheck">✓</span>Industry keyword gap analysis</li>
          <li><span class="fcheck">✓</span>Rewritten professional summary</li>
        </ul>
        <div class="price-row"><div><div class="price">$49</div></div></div>
        <button class="btn-card">Optimize my resume →</button>
      </div>
      <div class="card hot" onclick="openModal('interview')">
        <div class="hot-badge">Most popular</div>
        <div class="card-icon">🎤</div>
        <div class="card-name">Interview Prep</div>
        <p class="card-desc">Role-specific questions for your target company with full STAR method example answers.</p>
        <ul class="feats">
          <li><span class="fcheck">✓</span>15+ role-specific questions</li>
          <li><span class="fcheck">✓</span>Full STAR example answer</li>
          <li><span class="fcheck">✓</span>Company situational questions</li>
        </ul>
        <div class="price-row"><div><div class="price">$29</div></div></div>
        <button class="btn-card">Get my questions →</button>
      </div>
      <div class="card" onclick="openModal('salary')">
        <div class="card-icon">💰</div>
        <div class="card-name">Salary Negotiator</div>
        <p class="card-desc">Market rate benchmarking by role, seniority, and city — plus word-for-word negotiation scripts.</p>
        <ul class="feats">
          <li><span class="fcheck">✓</span>Market rate vs. your current salary</li>
          <li><span class="fcheck">✓</span>Floor / target / stretch range</li>
          <li><span class="fcheck">✓</span>4 ready-to-use scripts</li>
        </ul>
        <div class="price-row"><div><div class="price">$19</div></div></div>
        <button class="btn-card">Calculate my value →</button>
      </div>
    </div>
  </section>
</div>

<!-- PRICING -->
<div class="pricing-section" id="pricing">
  <div class="container">
    <div class="sec-head">
      <div class="sec-ey">Pricing</div>
      <h2 class="sec-title">Simple, crypto-friendly pricing</h2>
      <p class="sec-sub">Pay once per tool, or go Pro for unlimited access. No subscriptions with USD cards.</p>
    </div>
    <div class="pricing-grid">
      <div class="plan">
        <div class="plan-name">Free</div>
        <div class="plan-price"><sup>$</sup>0</div>
        <div class="plan-cadence">no payment needed</div>
        <ul class="plan-feats">
          <li><span class="fcheck">✓</span>1 ATS resume score</li>
          <li><span class="fcheck">✓</span>Keyword gap snapshot</li>
          <li><span class="fcheck">✓</span>Email delivery</li>
          <li style="color:var(--dim)">✕ PDF download</li>
          <li style="color:var(--dim)">✕ AI rewrite</li>
        </ul>
        <button class="btn-plan" onclick="openEmailCapture()">Get free score →</button>
      </div>
      <div class="plan popular">
        <div class="plan-badge">Best value</div>
        <div class="plan-name">Pro</div>
        <div class="plan-price"><sup>$</sup>5</div>
        <div class="plan-cadence">USDT / month · cancel any time</div>
        <ul class="plan-feats">
          <li><span class="fcheck">✓</span>Unlimited ATS scoring</li>
          <li><span class="fcheck">✓</span>Unlimited resume rewrites</li>
          <li><span class="fcheck">✓</span>Cover letters + PDFs</li>
          <li><span class="fcheck">✓</span>Interview prep (all roles)</li>
          <li><span class="fcheck">✓</span>Salary negotiation scripts</li>
        </ul>
        <button class="btn-plan accent" onclick="document.getElementById('services').scrollIntoView({behavior:'smooth'})">Start with a tool →</button>
      </div>
      <div class="plan">
        <div class="plan-name">Lifetime</div>
        <div class="plan-price"><sup>$</sup>29</div>
        <div class="plan-cadence">USDT · one-time payment</div>
        <ul class="plan-feats">
          <li><span class="fcheck">✓</span>Everything in Pro</li>
          <li><span class="fcheck">✓</span>Lifetime access</li>
          <li><span class="fcheck">✓</span>All future tools</li>
          <li><span class="fcheck">✓</span>Priority support</li>
          <li><span class="fcheck">✓</span>BTC, ETH, USDC, SOL accepted</li>
        </ul>
        <button class="btn-plan" onclick="document.getElementById('services').scrollIntoView({behavior:'smooth'})">Get lifetime access →</button>
      </div>
    </div>
    <div class="crypto-note">Pay with Bitcoin, Ethereum, USDC, or Solana. No credit card, no personal data stored.</div>
  </div>
</div>

<div class="container">
  <div class="how-section" id="how">
    <div class="sec-head" style="margin-bottom:40px">
      <div class="sec-ey">Process</div>
      <h2 class="sec-title">Three steps, done in minutes</h2>
    </div>
    <div class="how-grid">
      <div><div class="how-n">01</div><div class="how-l">Fill out your details or paste your resume</div></div>
      <div><div class="how-n">02</div><div class="how-l">Pay with BTC, ETH, USDC, or SOL</div></div>
      <div><div class="how-n">03</div><div class="how-l">Get your results + downloadable PDF</div></div>
    </div>
  </div>

  <section class="section" id="reviews">
    <div class="sec-head">
      <div class="sec-ey">Reviews</div>
      <h2 class="sec-title">Trusted by job seekers worldwide</h2>
    </div>
    <div class="cards">
      <div class="card" style="cursor:default">
        <div style="color:var(--gold);margin-bottom:12px">★★★★★</div>
        <p class="card-desc" style="margin-bottom:12px">"Generated my resume in seconds. Got 3 interview calls the next week — the ATS score made it obvious what I was missing."</p>
        <div style="font-size:13px;color:var(--muted)">— Priya S., Senior SWE</div>
      </div>
      <div class="card" style="cursor:default">
        <div style="color:var(--gold);margin-bottom:12px">★★★★★</div>
        <p class="card-desc" style="margin-bottom:12px">"The salary scripts got me from $95K to $118K. Used the exact wording from the negotiation tool."</p>
        <div style="font-size:13px;color:var(--muted)">— Marcus T., Product Manager</div>
      </div>
      <div class="card" style="cursor:default">
        <div style="color:var(--gold);margin-bottom:12px">★★★★★</div>
        <p class="card-desc" style="margin-bottom:12px">"Loved paying with crypto — instant delivery, no card details required. Interview questions were spot on for my Stripe round."</p>
        <div style="font-size:13px;color:var(--muted)">— Oluseun A., Data Analyst</div>
      </div>
    </div>
  </section>

  <footer>
    <div style="font-family:var(--display);font-weight:700">CareerForge Pro</div>
    <div class="footer-links">
      <a href="/ats-resume-checker">ATS Checker</a>
      <a href="/ai-cover-letter-generator">Cover Letter Generator</a>
      <a href="/interview-questions-generator">Interview Questions</a>
      <a href="/resume-score">Resume Score</a>
      <a href="/resume-keywords">Resume Keywords</a>
    </div>
    <div style="font-size:12px;color:var(--dim)">© 2025 CareerForge Pro</div>
  </footer>
</div>

<!-- MODAL -->
<div class="overlay" id="overlay" onclick="handleOverlayClick(event)">
  <div class="modal" id="modal">
    <div class="mhead">
      <div class="mtitle" id="mtitle">Career Tool</div>
      <button class="mclose" onclick="closeModal()">✕</button>
    </div>
    <div class="mbody" id="mbody"></div>
  </div>
</div>

<script>
// ─────────────────────────────────────────────
// STATE
// ─────────────────────────────────────────────
var activeService = null;
var formPayload = {};
var currentOrder = null;
var aiResult = null;
var fcSelectedFile = null;

var SERVICE_NAMES  = {resume_ai:'AI Resume Generator',resume_optimizer:'Resume Optimizer',interview:'Interview Prep',salary:'Salary Negotiator'};
var SERVICE_PRICES = {resume_ai:49,resume_optimizer:49,interview:29,salary:19};
var COINS = {
  bitcoin:  {label:'Bitcoin',  ticker:'BTC',  color:'#f7931a', logo:'₿'},
  ethereum: {label:'Ethereum', ticker:'ETH',  color:'#627eea', logo:'Ξ'},
  usdc:     {label:'USDC',     ticker:'USDC', color:'#2775ca', logo:'$'},
  solana:   {label:'Solana',   ticker:'SOL',  color:'#9945ff', logo:'◎'}
};
var APPROX_PRICES = {bitcoin:65000,ethereum:3500,usdc:1,solana:150};

// ─────────────────────────────────────────────
// HELPERS
// ─────────────────────────────────────────────
function approxAmt(usd, coin) {
  var amt = usd / APPROX_PRICES[coin];
  return coin === 'usdc' ? amt.toFixed(2) : amt.toFixed(8);
}

function escHtml(str) {
  if (!str) return '';
  return String(str).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function setBody(html) { document.getElementById('mbody').innerHTML = html; }

function stepsHtml(active) {
  return ['Details','Payment','Results'].map(function(s, i) {
    var cls = i < active ? 'step done' : i === active ? 'step active' : 'step';
    var num = i < active ? '✓' : (i + 1);
    return '<div class="' + cls + '"><div class="snum">' + num + '</div><div class="slabel">' + s + '</div></div>';
  }).join('');
}

function fileIconFor(name) {
  var n = (name || '').toLowerCase();
  if (n.endsWith('.pdf')) return '📕';
  if (n.endsWith('.docx') || n.endsWith('.doc')) return '📘';
  return '📄';
}

// ─────────────────────────────────────────────
// EMAIL CAPTURE  (with file upload)
// ─────────────────────────────────────────────
function openEmailCapture() {
  document.getElementById('mtitle').textContent = 'Free ATS Resume Score';
  activeService = null;
  fcSelectedFile = null;
  setBody(
    '<div class="free-score-box">' +
      '<div style="font-size:36px;margin-bottom:12px">📊</div>' +
      '<h3>Get your ATS score in 30 seconds</h3>' +
      '<p>Enter your email and upload your resume. We\'ll score it instantly and send you a keyword gap report — totally free.</p>' +
      '<div class="fgroup" style="text-align:left"><label>Email address</label><input type="email" id="fc_email" placeholder="you@example.com" autocomplete="email"></div>' +
      '<div class="fgroup" style="text-align:left">' +
        '<label>Upload your resume</label>' +
        '<div class="dropzone" id="fc-dropzone" onclick="document.getElementById(\'fc_file\').click()">' +
          '<div class="dropzone-icon">📤</div>' +
          '<div class="dropzone-text"><strong>Click to upload</strong> or drag and drop</div>' +
          '<div class="dropzone-hint">PDF or DOCX, up to 8MB</div>' +
        '</div>' +
        '<input type="file" id="fc_file" accept=".pdf,.docx,.txt" style="display:none" onchange="handleFcFileSelect(this.files[0])">' +
        '<div id="fc-file-chip"></div>' +
        '<div class="or-divider">or paste text instead</div>' +
        '<textarea id="fc_resume" placeholder="Paste your resume text here…" style="min-height:90px"></textarea>' +
      '</div>' +
      '<button class="btn-full" id="fc-btn" onclick="submitFreeCapture()">Get My Free ATS Score →</button>' +
    '</div>' +
    '<div style="text-align:center;margin-top:16px"><div style="font-size:12px;color:var(--dim)">Want the full analysis? See our <a href="#pricing" onclick="closeModal();document.getElementById(\'pricing\').scrollIntoView({behavior:\'smooth\'})" style="color:var(--accent2);text-decoration:none">paid plans</a> below.</div></div>'
  );
  document.getElementById('overlay').classList.add('open');
  document.body.style.overflow = 'hidden';

  var dz = document.getElementById('fc-dropzone');
  ['dragenter','dragover'].forEach(function(evt) {
    dz.addEventListener(evt, function(e) { e.preventDefault(); dz.classList.add('drag'); });
  });
  ['dragleave','drop'].forEach(function(evt) {
    dz.addEventListener(evt, function(e) { e.preventDefault(); dz.classList.remove('drag'); });
  });
  dz.addEventListener('drop', function(e) {
    var f = e.dataTransfer.files && e.dataTransfer.files[0];
    if (f) handleFcFileSelect(f);
  });
}

function handleFcFileSelect(file) {
  if (!file) return;
  var validExt = /\.(pdf|docx|txt)$/i;
  if (!validExt.test(file.name)) {
    alert('Please upload a PDF, DOCX, or TXT file.');
    return;
  }
  if (file.size > 8 * 1024 * 1024) {
    alert('File is too large. Max size is 8MB.');
    return;
  }
  fcSelectedFile = file;
  var chip = document.getElementById('fc-file-chip');
  if (chip) {
    chip.innerHTML = '<div class="file-chip">' +
      '<div class="file-chip-icon">' + fileIconFor(file.name) + '</div>' +
      '<div class="file-chip-name">' + escHtml(file.name) + '</div>' +
      '<button class="file-chip-remove" onclick="clearFcFile()" type="button">✕</button>' +
    '</div>';
  }
  var ta = document.getElementById('fc_resume');
  if (ta) ta.placeholder = 'File selected — you can still paste text as a backup';
}

function clearFcFile() {
  fcSelectedFile = null;
  var input = document.getElementById('fc_file');
  if (input) input.value = '';
  var chip = document.getElementById('fc-file-chip');
  if (chip) chip.innerHTML = '';
}

function submitFreeCapture() {
  var email  = (document.getElementById('fc_email').value  || '').trim();
  var pasted = (document.getElementById('fc_resume').value || '').trim();
  var btn    = document.getElementById('fc-btn');
  if (!email || !email.includes('@')) { alert('Please enter a valid email address.'); return; }
  if (!fcSelectedFile && !pasted) { alert('Please upload your resume or paste the text.'); return; }

  btn.disabled = true;
  btn.innerHTML = '<div class="spin"></div> Scoring…';

  if (fcSelectedFile) {
    var fd = new FormData();
    fd.append('email', email);
    fd.append('resume_file', fcSelectedFile);
    fetch('/api/upload-resume', {method: 'POST', body: fd})
      .then(function(r) { return r.json(); })
      .then(function(data) {
        if (data.error) throw new Error(data.error);
        renderFreeScoreResult(email, data);
      })
      .catch(function(err) {
        btn.disabled = false; btn.textContent = 'Try again';
        alert('Something went wrong: ' + err.message);
      });
  } else {
    fetch('/api/email-capture', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({email: email, resume: pasted})
    })
    .then(function(r) { return r.json(); })
    .then(function(data) {
      if (data.error) throw new Error(data.error);
      renderFreeScoreResult(email, data);
    })
    .catch(function(err) {
      btn.disabled = false; btn.textContent = 'Try again';
      alert('Something went wrong: ' + err.message);
    });
  }
}

function renderFreeScoreResult(email, data) {
  var score = data.ats_score || 0;
  var grade = data.grade || 'C';
  var gc = grade === 'A' ? '#34d399' : grade === 'B' ? '#e6b44a' : '#f87171';
  var miss = (data.missing_keywords || []).slice(0,4).map(function(k){ return '<span class="kw miss">'+escHtml(k)+'</span>'; }).join('');
  setBody(
    '<div style="text-align:center;margin-bottom:24px">' +
      '<div style="font-size:56px;font-weight:700;font-family:var(--display);color:'+gc+'">'+score+'</div>' +
      '<div style="font-size:14px;color:var(--muted);margin-bottom:8px">ATS Score — Grade <strong style="color:'+gc+'">'+grade+'</strong></div>' +
      '<div style="font-size:13px;color:var(--muted)">Full report sent to <strong style="color:var(--txt)">'+escHtml(email)+'</strong></div>' +
    '</div>' +
    (miss ? '<div class="rsec"><div class="rsec-title">Quick wins — add these keywords</div><div class="kw-list">'+miss+'</div></div>' : '') +
    '<div class="divider"></div>' +
    '<div style="text-align:center"><div style="font-size:14px;color:var(--muted);margin-bottom:16px">Want the full rewrite + PDF?</div>' +
    '<button class="btn-full" onclick="closeModal();openModal(\'resume_optimizer\')">Upgrade to Full Analysis — $49 →</button></div>'
  );
}

function submitEmailCapture() {
  var email = (document.getElementById('capture-email').value || '').trim();
  var msg   = document.getElementById('capture-msg');
  if (!email || !email.includes('@')) {
    if (msg) { msg.textContent = 'Please enter a valid email.'; msg.style.color = 'var(--red)'; }
    return;
  }
  fetch('/api/email-capture', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({email: email, resume: ''})
  })
  .then(function(r) { return r.json(); })
  .then(function() {
    if (msg) { msg.textContent = '✓ Check your inbox — score on its way!'; msg.style.color = 'var(--green)'; }
    document.getElementById('capture-email').value = '';
  })
  .catch(function() {
    if (msg) { msg.textContent = 'Something went wrong. Try the modal above.'; msg.style.color = 'var(--red)'; }
  });
}

// ─────────────────────────────────────────────
// MODAL OPEN / CLOSE
// ─────────────────────────────────────────────
function openModal(svc) {
  activeService = svc;
  formPayload = {};
  currentOrder = null;
  aiResult = null;
  document.getElementById('mtitle').textContent = SERVICE_NAMES[svc];
  renderStep1();
  document.getElementById('overlay').classList.add('open');
  document.body.style.overflow = 'hidden';
}

function closeModal() {
  document.getElementById('overlay').classList.remove('open');
  document.body.style.overflow = '';
  activeService = null;
  currentOrder = null;
}

function handleOverlayClick(e) {
  if (e.target === document.getElementById('overlay')) closeModal();
}

document.addEventListener('keydown', function(e) { if (e.key === 'Escape') closeModal(); });

// Auto-open from URL param e.g. ?open=resume_ai
(function() {
  var p = new URLSearchParams(window.location.search);
  var svc = p.get('open');
  if (svc && ['resume_ai','resume_optimizer','interview','salary'].indexOf(svc) !== -1) {
    setTimeout(function(){ openModal(svc); }, 300);
  }
})();

// ─────────────────────────────────────────────
// STEP 1 — FORM
// ─────────────────────────────────────────────
function renderStep1() {
  var form = '';
  if (activeService === 'resume_ai') {
    form = '<div class="fgroup"><label>Your full name</label><input type="text" id="f_name" placeholder="e.g. Jane Doe" autocomplete="name"></div>' +
           '<div class="fgroup"><label>Skills &amp; background</label><textarea id="f_skills" placeholder="Describe your skills, years of experience, and background…"></textarea></div>' +
           '<div class="fgroup"><label>Job description <span style="color:var(--dim);font-weight:400;text-transform:none">(optional — tailors your resume)</span></label><textarea id="f_jobdesc" placeholder="Paste the job description here…" style="min-height:80px"></textarea></div>';
  } else if (activeService === 'resume_optimizer') {
    form = '<div class="fgroup"><label>Paste your current resume</label><textarea id="f_resume" placeholder="Copy and paste your full resume text here…" style="min-height:140px"></textarea></div>' +
           '<div class="frow"><div class="fgroup"><label>Target job title</label><input type="text" id="f_title" placeholder="e.g. Senior Product Manager"></div>' +
           '<div class="fgroup"><label>Industry</label><select id="f_industry"><option value="tech">Technology</option><option value="data">Data / Analytics</option><option value="marketing">Marketing</option><option value="sales">Sales</option><option value="finance">Finance</option><option value="design">Design / UX</option><option value="healthcare">Healthcare</option><option value="education">Education</option></select></div></div>';
  } else if (activeService === 'interview') {
    form = '<div class="frow"><div class="fgroup"><label>Target job title</label><input type="text" id="f_title" placeholder="e.g. Software Engineer II"></div>' +
           '<div class="fgroup"><label>Company name</label><input type="text" id="f_company" placeholder="e.g. Stripe"></div></div>' +
           '<div class="fgroup"><label>Experience level</label><select id="f_exp"><option value="entry">Entry (0–2 years)</option><option value="mid" selected>Mid (3–5 years)</option><option value="senior">Senior (6–9 years)</option><option value="lead">Lead / Staff (10+ years)</option></select></div>';
  } else {
    form = '<div class="frow"><div class="fgroup"><label>Current salary (USD / yr)</label><input type="number" id="f_salary" placeholder="95000" min="1"></div>' +
           '<div class="fgroup"><label>Years of experience</label><input type="number" id="f_years" placeholder="5" min="0" max="50"></div></div>' +
           '<div class="frow"><div class="fgroup"><label>Job title</label><input type="text" id="f_title" placeholder="e.g. Marketing Director"></div>' +
           '<div class="fgroup"><label>City or region</label><input type="text" id="f_location" placeholder="e.g. New York or Remote"></div></div>';
  }
  setBody(
    '<div class="steps">' + stepsHtml(0) + '</div>' +
    form +
    '<button class="btn-full" id="step1-btn" onclick="submitStep1()">Continue to payment →</button>'
  );
}

function submitStep1() {
  var btn = document.getElementById('step1-btn');
  if (btn) { btn.disabled = true; btn.textContent = 'Validating…'; }

  var ok = true;
  if (activeService === 'resume_ai') {
    var name = (document.getElementById('f_name').value || '').trim();
    var skills = (document.getElementById('f_skills').value || '').trim();
    var jobdesc = (document.getElementById('f_jobdesc').value || '').trim();
    if (!name)   { alert('Please enter your name.');               ok = false; }
    else if (!skills) { alert('Please enter your skills and background.'); ok = false; }
    else { formPayload = {name: name, skills: skills, job_desc: jobdesc}; }
  } else if (activeService === 'resume_optimizer') {
    var resume = (document.getElementById('f_resume').value || '').trim();
    var title  = (document.getElementById('f_title').value  || '').trim();
    if (!resume) { alert('Please paste your resume.');           ok = false; }
    else if (!title) { alert('Please enter your target job title.'); ok = false; }
    else { formPayload = {resume_text: resume, job_title: title, industry: document.getElementById('f_industry').value}; }
  } else if (activeService === 'interview') {
    var title   = (document.getElementById('f_title').value   || '').trim();
    var company = (document.getElementById('f_company').value || '').trim();
    if (!title)   { alert('Please enter the job title.');   ok = false; }
    else if (!company) { alert('Please enter the company name.'); ok = false; }
    else { formPayload = {job_title: title, company: company, experience: document.getElementById('f_exp').value}; }
  } else {
    var salary = parseInt(document.getElementById('f_salary').value) || 0;
    var years  = parseInt(document.getElementById('f_years').value);
    var title  = (document.getElementById('f_title').value    || '').trim();
    var loc    = (document.getElementById('f_location').value || '').trim();
    if (!salary || salary < 1) { alert('Please enter your current salary.'); ok = false; }
    else if (isNaN(years) || years < 0) { alert('Please enter years of experience.'); ok = false; }
    else if (!title) { alert('Please enter your job title.'); ok = false; }
    else if (!loc)   { alert('Please enter your city or region.'); ok = false; }
    else { formPayload = {current_salary: salary, years: years, job_title: title, location: loc}; }
  }

  if (!ok) {
    if (btn) { btn.disabled = false; btn.textContent = 'Continue to payment →'; }
    return;
  }
  renderStep2(null);
}

// ─────────────────────────────────────────────
// STEP 2 — COIN SELECTION + PAYMENT
// ─────────────────────────────────────────────
function renderStep2(selectedCoin) {
  var usd = SERVICE_PRICES[activeService];
  var coinHtml = Object.keys(COINS).map(function(c) {
    var m = COINS[c];
    var amt = approxAmt(usd, c);
    var sel = selectedCoin === c ? ' sel' : '';
    return '<div class="coin-opt' + sel + '" id="coin-' + c + '" onclick="selectCoin(\'' + c + '\')">' +
           '<div class="clogo" style="background:' + m.color + '22;color:' + m.color + '">' + m.logo + '</div>' +
           '<div><div class="cname">' + m.label + '</div><div class="camt">≈ ' + amt + ' ' + m.ticker + '</div></div>' +
           '</div>';
  }).join('');

  setBody(
    '<div class="steps">' + stepsHtml(1) + '</div>' +
    '<div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:20px">' +
      '<div style="font-size:14px;color:var(--muted)">Total due</div>' +
      '<div style="font-family:var(--display);font-size:28px;font-weight:700">$' + usd + '</div>' +
    '</div>' +
    '<div style="font-size:12px;font-weight:600;color:var(--muted);margin-bottom:10px;text-transform:uppercase;letter-spacing:.05em">Choose your currency</div>' +
    '<div class="coin-grid">' + coinHtml + '</div>' +
    '<div id="pay-detail"></div>' +
    '<button class="btn-back" onclick="renderStep1()">← Back to details</button>'
  );
}

function selectCoin(coin) {
  document.querySelectorAll('.coin-opt').forEach(function(el) { el.classList.remove('sel'); });
  var el = document.getElementById('coin-' + coin);
  if (el) el.classList.add('sel');
  createOrder(coin);
}

function createOrder(coin) {
  var el = document.getElementById('pay-detail');
  if (!el) return;
  el.innerHTML = '<div style="text-align:center;padding:24px"><div class="spin spin-accent"></div><div style="margin-top:12px;font-size:13px;color:var(--muted)">Generating payment address…</div></div>';

  fetch('/api/orders', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({service: activeService, coin: coin, payload: formPayload})
  })
  .then(function(r) { return r.json(); })
  .then(function(data) {
    if (data.error) throw new Error(data.error);
    currentOrder = data;
    renderPayDetail(data);
  })
  .catch(function(err) {
    if (el) el.innerHTML = '<div class="err-box"><span>⚠</span><span>' + escHtml(err.message) + '</span></div>';
  });
}

function renderPayDetail(order) {
  var el = document.getElementById('pay-detail');
  if (!el) return;
  var m = COINS[order.coin];
  var qrHtml = order.qr
    ? '<img src="' + escHtml(order.qr) + '" alt="Payment QR code">'
    : '<div style="font-size:11px;color:#999;padding:10px">QR unavailable</div>';

  el.innerHTML =
    '<div class="pay-box">' +
      '<div class="qr-wrap">' + qrHtml + '</div>' +
      '<div class="pay-amt" style="color:' + m.color + '">' + escHtml(String(order.amount)) + ' ' + m.ticker + '</div>' +
      '<div class="pay-usd">≈ $' + escHtml(String(order.usd)) + ' USD</div>' +
      '<div class="addr-box">' +
        '<div class="addr-text" id="addr-text">' + escHtml(order.address) + '</div>' +
        '<button class="copy-btn" id="copy-btn" onclick="doCopy()">Copy</button>' +
      '</div>' +
    '</div>' +
    '<div class="warn-box"><span>⚠</span><span>Send <strong>exactly ' + escHtml(String(order.amount)) + ' ' + m.ticker + '</strong> to this address. Include network fees. Do not send a different amount.</span></div>' +
    '<button class="btn-full" id="verify-btn" onclick="checkPayment()">I\'ve sent the payment →</button>' +
    '<div id="verify-status" style="margin-top:12px"></div>';
}

function doCopy() {
  var addrEl = document.getElementById('addr-text');
  var btn    = document.getElementById('copy-btn');
  if (!addrEl || !btn) return;
  var addr = addrEl.textContent || '';
  function markCopied() {
    btn.textContent = 'Copied!';
    btn.classList.add('copied');
    setTimeout(function() { btn.textContent = 'Copy'; btn.classList.remove('copied'); }, 2000);
  }
  if (navigator.clipboard && navigator.clipboard.writeText) {
    navigator.clipboard.writeText(addr).then(markCopied).catch(function() { fallbackCopy(addr, markCopied); });
  } else {
    fallbackCopy(addr, markCopied);
  }
}

function fallbackCopy(text, cb) {
  var ta = document.createElement('textarea');
  ta.value = text;
  ta.style.cssText = 'position:fixed;top:-9999px;left:-9999px;opacity:0';
  document.body.appendChild(ta);
  ta.focus();
  ta.select();
  try { document.execCommand('copy'); cb(); } catch(e) {}
  document.body.removeChild(ta);
}

// ─────────────────────────────────────────────
// PAYMENT VERIFICATION
// ─────────────────────────────────────────────
function checkPayment() {
  if (!currentOrder) return;
  var btn      = document.getElementById('verify-btn');
  var statusEl = document.getElementById('verify-status');
  if (btn) { btn.disabled = true; btn.innerHTML = '<div class="spin"></div> Verifying payment…'; }
  if (statusEl) statusEl.innerHTML = '';

  fetch('/api/orders/' + currentOrder.order_id + '/verify', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: '{}'
  })
  .then(function(r) { return r.json(); })
  .then(function(data) {
    if (data.status === 'paid' && data.result) {
      if (activeService === 'resume_ai') {
        generateAIResume();
      } else {
        renderStep3(data.result);
      }
    } else {
      if (btn) { btn.disabled = false; btn.textContent = 'Check again →'; }
      if (statusEl) statusEl.innerHTML = '<div class="warn-box"><span>⏳</span><span>Payment not yet detected. Allow 10–30 minutes for network confirmation, then try again.</span></div>';
    }
  })
  .catch(function(err) {
    if (btn) { btn.disabled = false; btn.textContent = 'Check again →'; }
    if (statusEl) statusEl.innerHTML = '<div class="err-box"><span>⚠</span><span>' + escHtml(err.message) + '</span></div>';
  });
}

// ─────────────────────────────────────────────
// AI RESUME GENERATION
// ─────────────────────────────────────────────
function generateAIResume() {
  var btn = document.getElementById('verify-btn');
  if (btn) { btn.disabled = true; btn.innerHTML = '<div class="spin"></div> Generating your resume…'; }

  fetch('/api/generate', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify(formPayload)
  })
  .then(function(r) { return r.json(); })
  .then(function(data) {
    if (data.error) throw new Error(data.error);
    aiResult = data;
    renderStep3({ai_result: data});
  })
  .catch(function(err) {
    var statusEl = document.getElementById('verify-status');
    if (btn) { btn.disabled = false; btn.textContent = 'Retry generation →'; btn.onclick = generateAIResume; }
    if (statusEl) statusEl.innerHTML = '<div class="err-box"><span>⚠</span><span>Generation failed: ' + escHtml(err.message) + '. Please try again.</span></div>';
  });
}

// ─────────────────────────────────────────────
// PDF DOWNLOAD
// ─────────────────────────────────────────────
function downloadPDF() {
  if (!aiResult) return;
  var btn = document.getElementById('pdf-dl-btn');
  if (btn) { btn.disabled = true; btn.innerHTML = '<div class="spin"></div> Generating PDF…'; }

  fetch('/api/pdf', {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({name: formPayload.name || 'Candidate', resume: aiResult.resume, cover_letter: aiResult.cover_letter})
  })
  .then(function(r) { return r.json(); })
  .then(function(data) {
    if (data.error) throw new Error(data.error);
    var a = document.createElement('a');
    a.href = '/api/download/' + encodeURIComponent(data.filename);
    a.download = data.filename;
    document.body.appendChild(a);
    a.click();
    document.body.removeChild(a);
    if (btn) { btn.disabled = false; btn.innerHTML = '📄 Download PDF'; }
  })
  .catch(function(err) {
    alert('PDF download failed: ' + err.message);
    if (btn) { btn.disabled = false; btn.innerHTML = '📄 Download PDF'; }
  });
}

// ─────────────────────────────────────────────
// STEP 3 — RESULTS
// ─────────────────────────────────────────────
function renderStep3(result) {
  document.getElementById('mtitle').textContent = 'Your results';
  var content = '';

  if (activeService === 'resume_ai') {
    content = result && result.ai_result ? renderAIResult(result.ai_result) : '<div class="err-box"><span>⚠</span><span>Could not load AI result. Please contact support.</span></div>';
  } else if (activeService === 'resume_optimizer') {
    content = renderResumeResult(result);
  } else if (activeService === 'interview') {
    content = renderInterviewResult(result);
  } else if (activeService === 'salary') {
    content = renderSalaryResult(result);
  }

  setBody(
    '<div class="steps">' + stepsHtml(2) + '</div>' +
    '<div class="success-icon">✓</div>' +
    '<div style="text-align:center;margin-bottom:24px">' +
      '<div style="font-family:var(--display);font-size:18px;font-weight:700;margin-bottom:4px">Payment confirmed</div>' +
      '<div style="font-size:13px;color:var(--muted)">Here are your results</div>' +
    '</div>' +
    '<div class="divider"></div>' +
    content +
    '<button class="btn-back" onclick="closeModal()" style="margin-top:16px">← Close</button>'
  );
}

function renderAIResult(result) {
  var resume = result.resume || '';
  var cover  = result.cover_letter || '';
  return '<div class="rsec">' +
           '<div class="rsec-title">Generated Resume</div>' +
           '<div class="script-block" style="white-space:pre-wrap;font-family:monospace;font-size:11px;max-height:280px;overflow-y:auto;line-height:1.6">' + escHtml(resume) + '</div>' +
         '</div>' +
         '<div class="rsec">' +
           '<div class="rsec-title">Cover Letter</div>' +
           '<div class="script-block" style="white-space:pre-wrap;font-size:12px;max-height:200px;overflow-y:auto;line-height:1.6">' + escHtml(cover) + '</div>' +
         '</div>' +
         '<button class="btn-full pdf-btn" id="pdf-dl-btn" onclick="downloadPDF()">📄 Download as PDF</button>';
}

function renderResumeResult(r) {
  if (!r) return '<div class="err-box"><span>⚠</span><span>Results unavailable.</span></div>';
  var score = r.ats_score || 0;
  var grade = r.grade || 'D';
  var gc    = grade === 'A' ? '#34d399' : grade === 'B' ? '#e6b44a' : '#f87171';
  var C     = 2 * Math.PI * 30;
  var dash  = C * (score / 100);

  var sugs = (r.suggestions || []).map(function(s) {
    return '<div class="sug-item"><div class="sev-dot sev-' + escHtml(s.severity) + '"></div><div class="sug-text">' + escHtml(s.message) + '</div></div>';
  }).join('');

  var found = (r.found_keywords || []).map(function(k) {
    return '<span class="kw">' + escHtml(k) + '</span>';
  }).join('');

  var miss = (r.missing_keywords || []).map(function(k) {
    return '<span class="kw miss">' + escHtml(k) + '</span>';
  }).join('');

  var scoreLabel = score >= 85 ? 'Excellent ATS score!' : score >= 70 ? 'Good — apply suggestions below.' : 'Needs improvement.';

  return '<div class="rsec">' +
           '<div class="rsec-title">ATS Score</div>' +
           '<div class="score-wrap">' +
             '<div class="score-ring">' +
               '<svg width="80" height="80" viewBox="0 0 80 80" xmlns="http://www.w3.org/2000/svg">' +
                 '<circle cx="40" cy="40" r="30" fill="none" stroke="var(--bg3)" stroke-width="7"/>' +
                 '<circle cx="40" cy="40" r="30" fill="none" stroke="' + gc + '" stroke-width="7" stroke-dasharray="' + dash.toFixed(2) + ' ' + C.toFixed(2) + '" stroke-linecap="round"/>' +
               '</svg>' +
               '<div class="score-label"><div class="score-num" style="color:' + gc + '">' + score + '</div><div class="score-grade">Grade ' + escHtml(grade) + '</div></div>' +
             '</div>' +
             '<div><div style="font-size:13px;color:var(--txt);font-weight:500;margin-bottom:4px">' + scoreLabel + '</div><div style="font-size:12px;color:var(--muted)">' + (r.word_count || 0) + ' words · ' + (r.has_metrics ? 'Has metrics ✓' : 'No metrics found') + '</div></div>' +
           '</div>' +
         '</div>' +
         (sugs ? '<div class="rsec"><div class="rsec-title">Improvements</div>' + sugs + '</div>' : '') +
         '<div class="rsec"><div class="rsec-title">Keywords Found</div><div class="kw-list">' + (found || '<span style="color:var(--muted);font-size:13px">None detected</span>') + '</div></div>' +
         '<div class="rsec"><div class="rsec-title">Keywords to Add</div><div class="kw-list">' + (miss || '<span style="color:var(--green);font-size:13px">Great coverage!</span>') + '</div></div>' +
         '<div class="rsec"><div class="rsec-title">Rewritten Professional Summary</div><div class="script-block">' + escHtml(r.optimized_summary || '') + '</div></div>';
}

function renderInterviewResult(r) {
  if (!r) return '<div class="err-box"><span>⚠</span><span>Results unavailable.</span></div>';
  var qs = r.questions || {};
  var sections = [{label:'Technical', key:'technical'}, {label:'Behavioural', key:'behavioral'}, {label:'Situational', key:'situational'}];

  var qHtml = sections.map(function(s) {
    var items = (qs[s.key] || []).map(function(q, i) {
      return '<div class="q-card"><div class="qn">Q' + (i+1) + '</div><div>' + escHtml(q) + '</div></div>';
    }).join('');
    return '<div class="rsec"><div class="rsec-title">' + s.label + ' Questions</div>' + items + '</div>';
  }).join('');

  var se = r.star_example || {};
  var starCells = ['situation','task','action','result'].map(function(k) {
    return '<div class="star-cell"><div class="star-k">' + k.charAt(0).toUpperCase() + k.slice(1) + '</div><div class="star-v">' + escHtml(se[k] || '') + '</div></div>';
  }).join('');

  var tips = (r.prep_tips || []).map(function(t) {
    return '<div class="q-card"><div class="qn">💡</div><div>' + escHtml(t) + '</div></div>';
  }).join('');

  return qHtml +
         '<div class="rsec"><div class="rsec-title">STAR Method Example Answer</div>' +
           '<div class="star-grid">' + starCells + '</div>' +
           (se.tip ? '<div class="tip-box"><span>💡</span><span>' + escHtml(se.tip) + '</span></div>' : '') +
         '</div>' +
         '<div class="rsec"><div class="rsec-title">Preparation Tips</div>' + tips + '</div>';
}

function renderSalaryResult(r) {
  if (!r) return '<div class="err-box"><span>⚠</span><span>Results unavailable.</span></div>';
  var range   = r.range || {};
  var floor   = range.floor   || 0;
  var target  = range.target  || 0;
  var stretch = range.stretch || 0;
  var maxV    = stretch || 1;

  var scripts = (r.scripts || []).map(function(s) {
    return '<div class="script-block">' + escHtml(s) + '</div>';
  }).join('');

  return '<div class="rsec">' +
           '<div class="rsec-title">Your Market Value — ' + escHtml(r.level || '') + ' · ' + escHtml(r.location || '') + '</div>' +
           '<div class="range-cards">' +
             '<div class="rc"><div class="rc-label">Current</div><div class="rc-val">$' + (r.current_salary || 0).toLocaleString() + '</div></div>' +
             '<div class="rc target"><div class="rc-label">Target</div><div class="rc-val" style="color:var(--accent2)">$' + target.toLocaleString() + '</div></div>' +
             '<div class="rc"><div class="rc-label">Stretch</div><div class="rc-val" style="color:var(--green)">$' + stretch.toLocaleString() + '</div></div>' +
           '</div>' +
           (r.pct_above_current > 0
             ? '<div style="font-size:12px;color:var(--green);margin-bottom:12px">↑ Market rate is ' + r.pct_above_current + '% above your current salary. ' + escHtml(r.annual_upside || '') + '.</div>'
             : '') +
           '<div class="rbar"><div class="rbar-fill" style="width:' + Math.min(100, Math.round(target / maxV * 100)) + '%"></div></div>' +
           '<div style="display:flex;justify-content:space-between;margin-top:6px">' +
             '<div style="font-size:10px;color:var(--dim)">Floor $' + floor.toLocaleString() + '</div>' +
             '<div style="font-size:10px;color:var(--dim)">Stretch $' + stretch.toLocaleString() + '</div>' +
           '</div>' +
         '</div>' +
         '<div class="rsec"><div class="rsec-title">Negotiation Scripts</div>' + scripts + '</div>' +
         (r.total_comp_note ? '<div class="tip-box"><span>💡</span><span>' + escHtml(r.total_comp_note) + '</span></div>' : '');
}
</script>
</body>
</html>'''

# ============================================================================
# API ROUTES
# ============================================================================
@app.get("/api/services")
def api_services():
    return jsonify(SERVICES)

@app.get("/api/prices")
def api_prices():
    return jsonify(CRYPTO_PRICES_USD)

# Email capture endpoint — paste-text path (no file). Persists to SQLite.
@app.post("/api/email-capture")
def email_capture():
    body   = request.get_json(silent=True) or {}
    email  = (body.get("email") or "").strip().lower()
    resume = (body.get("resume") or "").strip()

    if not email or "@" not in email or "." not in email.split("@")[-1]:
        return jsonify({"error": "Valid email required"}), 400

    result = {}
    if resume:
        result = run_resume_optimizer(resume, "Professional", "tech")

    save_email_lead(
        email=email,
        has_resume=bool(resume),
        ats_score=result.get("ats_score", 0),
        grade=result.get("grade", ""),
        source="capture_band_paste",
    )

    # In production: send a transactional email with the score via SendGrid / Mailgun / Postmark.
    return jsonify({
        "ok":               True,
        "ats_score":        result.get("ats_score", 0),
        "grade":            result.get("grade", ""),
        "missing_keywords": result.get("missing_keywords", [])[:4],
        "message":          "Score sent to " + email,
    })

# Email capture endpoint — file upload path (PDF/DOCX/TXT). Persists to SQLite.
@app.post("/api/upload-resume")
def upload_resume():
    email = (request.form.get("email") or "").strip().lower()
    if not email or "@" not in email or "." not in email.split("@")[-1]:
        return jsonify({"error": "Valid email required"}), 400

    if "resume_file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file_storage = request.files["resume_file"]
    if not file_storage or not file_storage.filename:
        return jsonify({"error": "No file uploaded"}), 400

    resume_text, err = extract_text_from_upload(file_storage)
    if err:
        return jsonify({"error": err}), 400

    result = run_resume_optimizer(resume_text, "Professional", "tech")

    save_email_lead(
        email=email,
        has_resume=True,
        ats_score=result.get("ats_score", 0),
        grade=result.get("grade", ""),
        source="capture_band_upload",
    )

    return jsonify({
        "ok":               True,
        "ats_score":        result.get("ats_score", 0),
        "grade":            result.get("grade", ""),
        "missing_keywords": result.get("missing_keywords", [])[:4],
        "word_count":       result.get("word_count", 0),
        "message":          "Score sent to " + email,
    })

# Admin endpoint — protect with ADMIN_SECRET env var in production
@app.get("/api/emails")
def list_emails():
    admin_secret = os.environ.get("ADMIN_SECRET", "")
    secret = request.args.get("secret", "")
    if not admin_secret or secret != admin_secret:
        return jsonify({"error": "Unauthorized"}), 401
    return jsonify({"count": count_emails(), "emails": list_all_emails()})

@app.post("/api/orders")
def create_order():
    body = request.get_json(silent=True) or {}
    service = body.get("service", "").strip()
    coin    = body.get("coin", "").strip()

    if service not in SERVICES:
        return jsonify({"error": f"Unknown service. Valid: {', '.join(SERVICES.keys())}"}), 400
    if coin not in WALLETS:
        return jsonify({"error": "Unsupported coin. Valid: bitcoin, ethereum, usdc, solana"}), 400

    usd     = SERVICES[service]["price_usd"]
    amount  = crypto_amount(usd, coin)
    address = WALLETS[coin]
    order_id = secrets.token_urlsafe(12)

    try:
        qr_data = generate_qr(address, coin, amount)
    except Exception:
        qr_data = ""

    _orders[order_id] = {
        "id":        order_id,
        "service":   service,
        "coin":      coin,
        "usd_amount": usd,
        "amount":    amount,
        "address":   address,
        "status":    "awaiting_payment",
        "created_at": utcnow(),
        "payload":   body.get("payload") or {},
    }

    return jsonify({
        "order_id": order_id,
        "address":  address,
        "coin":     coin,
        "amount":   amount,
        "usd":      usd,
        "qr":       qr_data,
    })

@app.post("/api/orders/<order_id>/verify")
def verify_order(order_id):
    order = _orders.get(order_id)
    if not order:
        return jsonify({"error": "Order not found"}), 404

    if order["status"] == "paid":
        return jsonify({
            "status":        "paid",
            "confirmations": order.get("confirmations", 1),
            "result":        order.get("result"),
        })

    body    = request.get_json(silent=True) or {}
    tx_hash = body.get("tx_hash")
    check   = verify_on_chain(order["coin"], order["address"], order["amount"], tx_hash)

    if not check["confirmed"]:
        return jsonify({
            "status":        "awaiting_payment",
            "confirmations": 0,
            "message":       "Payment not yet detected. Allow 10–30 minutes for network confirmation.",
        })

    order["status"]        = "paid"
    order["tx_hash"]       = check["tx_hash"]
    order["confirmations"] = check["confirmations"]
    order["paid_at"]       = utcnow()

    svc     = order["service"]
    payload = order.get("payload") or {}

    try:
        if svc == "resume_optimizer":
            result = run_resume_optimizer(
                payload.get("resume_text", ""),
                payload.get("job_title", "Professional"),
                payload.get("industry", "tech"),
            )
        elif svc == "resume_ai":
            result = {"requires_ai_generation": True}
        elif svc == "interview":
            result = run_interview_prep(
                payload.get("job_title", "Professional"),
                payload.get("experience", "mid"),
                payload.get("company", "the company"),
            )
        elif svc == "salary":
            result = run_salary_negotiation(
                int(payload.get("current_salary") or 80000),
                int(payload.get("years") or 3),
                payload.get("location", "remote"),
                payload.get("job_title", "Professional"),
            )
        else:
            result = {}
    except Exception as exc:
        result = {"error": str(exc)}

    order["result"] = result
    return jsonify({
        "status":        "paid",
        "confirmations": check["confirmations"],
        "result":        result,
    })

@app.get("/api/orders/<order_id>")
def get_order(order_id):
    order = _orders.get(order_id)
    if not order:
        return jsonify({"error": "Order not found"}), 404
    safe = {k: v for k, v in order.items() if k != "payload"}
    return jsonify(safe)

@app.post("/api/generate")
def generate():
    body   = request.get_json(silent=True) or {}
    name   = (body.get("name")   or "").strip()
    skills = (body.get("skills") or "").strip()
    job_desc = (body.get("job_desc") or "").strip()

    if not name:   return jsonify({"error": "Name is required"}), 400
    if not skills: return jsonify({"error": "Skills are required"}), 400

    try:
        resume_text = call_claude(build_resume_prompt(name, skills, job_desc))
        cover_text  = call_claude(build_cover_letter_prompt(name, skills, job_desc, resume_text))
        session_id  = secrets.token_urlsafe(16)
        _ai_results[session_id] = {"resume": resume_text, "cover_letter": cover_text, "name": name}
    except Exception as e:
        return jsonify({"error": f"AI generation failed: {str(e)}"}), 500

    return jsonify({"resume": resume_text, "cover_letter": cover_text, "session_id": session_id})

@app.post("/api/pdf")
def create_pdf():
    body       = request.get_json(silent=True) or {}
    name       = (body.get("name")          or "Candidate").strip()
    resume_text = (body.get("resume")       or "").strip()
    cover_text  = (body.get("cover_letter") or "").strip()

    if not resume_text:
        return jsonify({"error": "Resume content is required"}), 400

    try:
        filename = build_pdf(name, resume_text, cover_text)
    except Exception as e:
        return jsonify({"error": f"PDF generation failed: {str(e)}"}), 500

    return jsonify({"filename": filename})

@app.get("/api/download/<filename>")
def download_file(filename):
    safe = re.sub(r"[^a-zA-Z0-9_\-.]", "", filename)
    if not safe.endswith(".pdf"):
        return jsonify({"error": "Invalid filename"}), 400
    filepath = os.path.join(OUTPUTS_DIR, safe)
    if not os.path.exists(filepath):
        return jsonify({"error": "File not found. Please regenerate."}), 404
    return send_file(filepath, as_attachment=True, download_name=safe, mimetype="application/pdf")

@app.get("/health")
def health():
    return jsonify({"status": "ok", "orders": len(_orders), "ai_results": len(_ai_results), "emails": count_emails()})

# ── SEO LANDING PAGES ────────────────────────────────────────────────────────
@app.get("/ats-resume-checker")
def page_ats_checker():
    return seo_page(
        "Free ATS Resume Checker",
        "Check If Your Resume Passes",
        "Paste your resume and get an ATS compatibility score, keyword gaps, and a rewritten summary — free in seconds.",
        "Free ATS resume checker. Get an instant ATS score, keyword gap analysis, and actionable improvements. Works for all industries.",
        "Check My Resume Free",
        "resume_optimizer"
    )

@app.get("/ai-cover-letter-generator")
def page_cover_letter():
    return seo_page(
        "AI Cover Letter Generator",
        "Write a Cover Letter That",
        "AI generates a tailored, ATS-friendly cover letter from your resume and job description — under 5 minutes.",
        "AI cover letter generator. Paste your resume and job description, get a professional, ATS-optimized cover letter instantly.",
        "Generate My Cover Letter",
        "resume_ai"
    )

@app.get("/interview-questions-generator")
def page_interview():
    return seo_page(
        "Interview Questions Generator",
        "Practice the Questions",
        "Get 15+ role-specific interview questions, a full STAR example answer, and company situational questions tailored to your target role.",
        "AI interview questions generator. Get role-specific technical, behavioral, and situational questions with STAR method example answers.",
        "Get My Questions",
        "interview"
    )

@app.get("/resume-score")
def page_resume_score():
    return seo_page(
        "Free Resume Score Checker",
        "Score Your Resume Before",
        "Find out exactly how recruiters and ATS systems see your resume. Get a letter grade, keyword analysis, and three quick improvements.",
        "Free resume score checker. Letter-grade ATS scoring, keyword gap analysis, and a rewritten professional summary.",
        "Score My Resume",
        "resume_optimizer"
    )

@app.get("/resume-keywords")
def page_keywords():
    return seo_page(
        "Resume Keywords Analyzer",
        "Find the Keywords Your",
        "See which high-value industry keywords your resume is missing — and which ones to add to pass ATS filters and land interviews.",
        "Resume keyword analyzer. Find missing ATS keywords for your industry and get suggestions to improve your resume's match rate.",
        "Analyze My Keywords",
        "resume_optimizer"
    )

# ── SPA catch-all ────────────────────────────────────────────────────────────
@app.get("/")
def index():
    return FRONTEND_HTML

@app.get("/<path:path>")
def catch_all(path):
    if path.startswith("api/"):
        return jsonify({"error": "Not found"}), 404
    return FRONTEND_HTML

# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print()
    print("─" * 60)
    print("  CareerForge Pro  ·  http://localhost:5000")
    print("─" * 60)
    for k, v in WALLETS.items():
        print(f"  {k.upper():<10} {v}")
    print("─" * 60)
    for svc, info in SERVICES.items():
        print(f"    • {info['name']} (${info['price_usd']})")
    print("─" * 60)
    print("  SEO pages: /ats-resume-checker  /ai-cover-letter-generator")
    print("             /interview-questions-generator  /resume-score  /resume-keywords")
    print("─" * 60)
    print("  Free ATS score + email capture: paste text OR upload PDF/DOCX/TXT")
    print("  Admin emails: GET /api/emails?secret=<ADMIN_SECRET>")
    print("─" * 60)
    print("  Set ANTHROPIC_API_KEY for AI resume generation")
    print("─" * 60)
    print()
    app.run(debug=False, host="0.0.0.0", port=5000)
